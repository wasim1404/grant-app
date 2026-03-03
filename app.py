import streamlit as st
import streamlit.components.v1 as components
import re
import io
import PyPDF2
import requests
from bs4 import BeautifulSoup
from readability import Document
import trafilatura
from urllib.parse import urljoin
import os
from pathlib import Path
import base64
import json
import csv
import sqlite3
import traceback
from fpdf import FPDF as PDF
from typing import Optional

# OpenAI client
from openai import OpenAI

# Conditional import for python-docx
try:
    import docx
except ImportError:
    docx = None # Set to None if not available
    # The error message for docx will be handled where it's used.

import sqlite3
from datetime import datetime, timedelta
import time

# --- Global Variables & Constants ---
BASE_DIR = Path(__file__).resolve().parent
DATABASE_FILE = str(BASE_DIR / "proposals.db")
# DST announcements (primary source for Grant Finder opportunities)
DST_ANNOUNCEMENTS_URL = "https://dst.gov.in/whatsnew/announcement"
# ANRF portal homepage (secondary source for Grant Finder opportunities)
ANRF_HOMEPAGE_URL = "https://www.anrfonline.in/ANRF/HomePage"
# Community-maintained opportunity sheet (public CSV export required)
OPPORTUNITY_SHEET_URL = "https://docs.google.com/spreadsheets/d/1QkJxZ-_mep-vgNLSfHhLr6ngKOi02Tj5QLD3FBqStzE/edit?pli=1&gid=1096084732#gid=1096084732"
OPPORTUNITY_SHEET_ID = "1QkJxZ-_mep-vgNLSfHhLr6ngKOi02Tj5QLD3FBqStzE"
OPPORTUNITY_SHEET_GID = "1096084732"
# India Science & Technology latest updates (additional source)
INDIA_SCI_TECH_LATEST_URL = "https://www.indiascienceandtechnology.gov.in/latest-updates"
# Fallback taxonomy used if taxonomy.json is missing in deployment
DEFAULT_TAXONOMY = {
    "Sustainable Development and Ecology": [
        "Climate Change Mitigation and Adaptation",
        "Integrated Water Resource Management",
        "Wildlife Conservation and Ecology",
        "Renewable Resource Management",
        "Circular Economy and Sustainable Business Practices",
        "Sustainable Agriculture and Food Security",
        "Urban Green Spaces and Biodiversity",
        "Environmental Education and Awareness",
        "Fisheries, Animal Husbandry and Dairy",
        "Land Restoration and Reforestation",
        "Waste to Wealth",
    ],
    "Health and Well-being": [
        "Healthcare Innovation and Technology",
        "Rural Health and Well Being",
        "Traditional Medicine (AYUSH)",
        "Cancer & Metabolic Disorders",
        "Drug Abuse & Addiction",
        "Epidemiology and Pandemic Response",
        "Global Health Disparities",
        "Elderly Health care and Wellness",
        "Cognitive Neuroscience and Brain Health",
        "Mental Health and Well-being",
        "Nutrition and Dietetics",
        "Personalized Medicine and Genomics",
        "Infectious Diseases and Vaccinology",
        "Specially abled (Divyang) Wellness",
        "Maternal and Child Health",
        "Drug Design and Discovery",
    ],
    "Technology and Innovation": [
        "AI Ethics and Society",
        "Quantum Computing and Applications",
        "Innovative Educational Technologies",
        "Human-Robot Interaction",
        "Virtual Reality and Immersive Experiences",
        "Internet of Things (IoT)",
        "Big Data Analytics and Applications",
        "Blockchain Technology",
        "Smart Cities, Road and Urban Innovation",
        "Renewable Energy Technologies",
        "5G,6G & beyond",
        "Empowerment of person with Disabilities",
        "AI Based Technology & Innovation",
        "Digital twin Technology",
        "Industry 4.0",
        "Agricultural Technology & Innovation",
    ],
    "Environmental Science and Energy": [
        "Alternative Energy Solutions",
        "Sustainable Materials Development",
        "Artificial Photosynthesis and Carbon Capture",
        "Eco-Architecture and Green Buildings",
        "Ocean Sciences and Coastal Studies",
        "Soil Science and Land Management",
        "Waste Management and Recycling",
        "Environmental Monitoring and Assessment",
        "Climate Change Impacts and Resilience",
        "Energy Storage and Battery Technology",
        "Conventional Energy Sources",
    ],
    "Arts, Media, and Culture": [
        "Digital Arts and Culture",
        "Music Technology and Sound Engineering",
        "Theatre and Digital Performance",
        "Cultural Heritage Management",
        "Creative Writing and Multimedia Content",
        "Visual Arts and Design",
        "Film and Media Studies",
        "Fashion and Textile Innovation",
        "Performing Arts and Dance",
        "Art Therapy and Community Arts",
    ],
    "Social Sciences and Global Studies": [
        "Socio-Economic Policy Analysis",
        "Civil Rights and Social Equity",
        "Feminist Studies and Gender Advocacy",
        "Labour and Employment Studies",
        "Refugee and Migration Studies",
        "Geopolitics and International Security",
        "Urban Studies and Planning",
        "Rural Development and Agrarian Studies",
        "Conflict Resolution and Peace Studies",
        "Public Policy and Governance",
        "Anthropology and Cultural Studies",
        "Constitutional, Legislative & Judicial reforms",
    ],
    "Business and Economics": [
        "Business Ethics and Corporate Governance",
        "Behavioural Economics and Decision Sciences",
        "Sports Performance and Management",
        "Corporate Social Responsibility and Ethics",
        "Sustainable Tourism and Heritage Conservation",
        "Innovation and Entrepreneurship",
        "Financial Technologies (FinTech)",
        "International Trade and Global Markets",
        "Marketing and Consumer Behavior",
        "Human Resource Management and Organizational Behavior",
        "Ports, Shipping & Waterways",
    ],
    "Security and Privacy": [
        "Cybersecurity and Data Privacy",
        "Privacy and Security in the Digital Age",
        "Cyber-Physical Systems Security",
        "Blockchain for Social Impact",
        "Community Resilience and Disaster Management",
        "Digital Forensics and Incident Response",
        "Network Security and Cryptography",
        "Information Assurance and Risk Management",
        "Biometric Security and Identity Management",
        "Critical Infrastructure Protection",
        "National Security",
    ],
    "Engineering and Technology": [
        "Space Exploration and Satellite Technologies",
        "Renewable Energy Systems",
        "Nanoengineering for Medicine and Electronics (VLSI etc)",
        "Robotics and Autonomous Systems",
        "Advanced Food Production & Packaging Techniques",
        "Sustainable Manufacturing Processes",
        "Smart Grid and Energy Management",
        "Advanced Transportation Systems",
        "Biomechanics and Biomedical Engineering",
        "Additive Manufacturing and 3D Printing",
        "Atomic Energy & Nuclear Technology",
        "Mine Technology",
        "Communication Technology",
        "Supply Chain Management Technology",
        "Applied Mathematics and Statistics",
        "Laser, Plasma and Matter",
    ],
    "Education and Community Engagement": [
        "Youth Development, Training and Leadership",
        "Rural Education",
        "Train the Trainer",
        "Educational Technologies and Pedagogies",
        "Accessibility and Inclusive Design",
        "Special Needs and Inclusive Education",
        "One Health and Veterinary Sciences",
        "Community Health and Well-being",
        "Lifelong Learning and Adult Education",
        "Civic Engagement and Emotional Education",
        "Teacher Training and Professional Development",
        "Family and Child Studies",
        "Demographic Studies",
        "Equitable & Inclusive Society",
        "Women Empowerment",
        "National Educational Policy 2020",
        "Sports Infrastructure & Innovation",
        "Employment Generation & Skill Development",
    ],
    "Interdisciplinary Research": [
        "Sustainable Development Goals (SDGs) Integration",
        "Technological Innovations for Global Challenges",
        "Ethical Implications of Emerging Technologies",
        "Interdisciplinary Approaches to Climate Change",
        "Human-Machine Collaboration and Augmentation",
    ],
    "Cultural and Heritage Studies": [
        "Digital Preservation of Cultural Heritage",
        "Traditional Knowledge and Indigenous Studies",
        "Urban Heritage and Regeneration",
        "Heritage Tourism and Community Development",
        "Cultural Diplomacy and International Relations",
        "Indigenous Innovation",
    ],
    "Economic Development and Policy": [
        "Economic Development in Emerging Markets",
        "Trade Policy and Economic Integration",
        "Microfinance and Social Entrepreneurship",
        "Industrial Policy and Economic Planning",
        "Public Finance and Fiscal Policy",
        "Start Up & Entrepreneurship Training",
    ],
    "Environmental Health and Safety": [
        "Occupational Health and Safety",
        "Environmental Health Risk Assessment",
        "Toxicology and Chemical Safety",
        "Public Health Infrastructure and Policy",
        "Disaster Management",
        "Radiation studies and Solutions",
        "Waste water treatment and assessment",
    ],
}
# OpenAI API Key
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
SELECTED_MODEL = "gpt-4o-mini"

# Initialize OpenAI client (graceful if key missing)
if OPENAI_API_KEY:
    openai_client = OpenAI(api_key=OPENAI_API_KEY)
else:
    openai_client = None
    st.error("OPENAI_API_KEY is not set. Please set it in your environment or Streamlit Secrets.")

# --- Rate Limiting ---
LAST_API_CALL_TIME = 0
MIN_SECONDS_BETWEEN_CALLS = 1  # OpenAI has generous rate limits

# --- Response wrapper to maintain compatibility ---
class AIResponse:
    def __init__(self, text):
        self.text = text

# --- AI Helper Functions ---
def generate_content_with_retry(model_name, prompt, max_retries=5, delay=5):
    global LAST_API_CALL_TIME
    if openai_client is None:
        st.error("OPENAI_API_KEY is missing. Set it and restart the app to use AI features.")
        return None
    
    # Throttle: ensure minimum time between API calls
    current_time = time.time()
    time_since_last_call = current_time - LAST_API_CALL_TIME
    if time_since_last_call < MIN_SECONDS_BETWEEN_CALLS:
        wait_needed = MIN_SECONDS_BETWEEN_CALLS - time_since_last_call
        time.sleep(wait_needed)
    
    for i in range(max_retries):
        try:
            LAST_API_CALL_TIME = time.time()
            response = openai_client.chat.completions.create(
                model=model_name,
                messages=[{"role": "user", "content": prompt}],
                max_tokens=4096
            )
            return AIResponse(response.choices[0].message.content)
        except Exception as e:
            error_str = str(e).lower()
            if "rate" in error_str or "limit" in error_str or "quota" in error_str:
                if i < max_retries - 1:
                    wait_time = delay * (2 ** i)
                    st.warning(f"Rate limit hit. Waiting {wait_time} seconds before retry {i+2}/{max_retries}...")
                    time.sleep(wait_time)
                else:
                    st.error("API rate limit exceeded. Please try again in a few minutes.")
                    return None
            else:
                st.error(f"An unexpected error occurred during content generation: {e}")
                traceback.print_exc()
                return None

def summarize_text_for_prompt(text, max_length=2000, model_name=SELECTED_MODEL):
    # Define a hard limit for the input to the summarization prompt itself
    SUMMARIZATION_INPUT_HARD_LIMIT = 5000 # Characters - to prevent summarization prompt from being too long

    if len(text) > max_length:
        # Truncate text for the summarization prompt if it's extremely long
        text_for_summarization = text[:SUMMARIZATION_INPUT_HARD_LIMIT]
        
        # Target summary length (e.g., 75% of max_length for the main prompt)
        target_summary_length = int(max_length * 0.75)

        summary_prompt = f"""
        Summarize the following text concisely, retaining all critical information for a research proposal context. The summary should be approximately {target_summary_length} characters long.

        Text to Summarize:
        {text_for_summarization}
        """
        st.info("Summarizing lengthy input to fit AI context window...")
        summary_response = generate_content_with_retry(model_name, summary_prompt)
        if summary_response and summary_response.text:
            # Ensure the summary itself doesn't exceed the intended max_length
            return summary_response.text[:max_length]
        else:
            st.warning("Failed to summarize text. Using original (potentially truncated) text.")
            return text[:max_length] # Fallback to truncation if summarization fails
    return text

def extract_fields(text):
    fields = {
        "Funding Agency": "N/A",
        "Scheme Type": "N/A",
        "Duration": "N/A",
        "Budget": "N/A",
        "Thrust Areas": "N/A",
        "Eligibility": "N/A",
        "Submission Format": "N/A",
        "Last Date of Submission": "N/A",
        "Scheme or Call Name": "N/A",
        "Scope or Objective of the Programme": "N/A",
    }

    # More robust regex patterns
    patterns = {
        "Funding Agency": r"(Indian Council of Medical Research \(ICMR\)|ICMR)",
        "Scheme Type": r"Scheme Type[s]?:\s*(.*?)(?:\n|$)",
        "Duration": r"(?:(?:maximum|minimum|flexible) duration.*?)((?:\d+\s*years)|flexible duration)",
        "Budget": r"(up to \d+\s*Cr each|funding will be linked to deliverables)",
        "Thrust Areas": r"(novel, futuristic ideas, new knowledge generation, discovery/ development of breakthrough health technologies)",
        "Eligibility": r"Eligibility[:\s]*([\s\S]*?)(?=\n\n|Application:|Selection|Timelines|\Z)",
        "Submission Format": r"(Proposal must be submitted only through e-PMS portal of ICMR)",
        "Last Date of Submission": r"(?:Last Date of Submission|Call is open until|Submission Deadline)[:\s]*(.*?)(?:\n|$)|(Rolling call and received proposals will be evaluated every month\.)",
        "Scheme or Call Name": r"(Call for R&D project Proposals|Scheme or Call Name[:\s]*(.*?)(?:\n|$))",
        "Scope or Objective of the Programme": r"(VISHLESHAN I-HUB FOUNDATION, IIT Patna is the nodal centre and a Technology Innovation Hub \(TIH\) for technology development and activities in the core areas of 'Speech, Video, and Text Analytics Technologies' in synergy with 'wireless, sensor, and IoT technologies, material sciences etc\. under National Mission on Interdisciplinary Cyber Physical Systems \(NM-ICPS\)|Problem Statements[\s\S]*?(?:Review:|Multimedia Networking|Multimedia Traffic Management|Image/Video Security and Privacy|Real-time image and video processing|Robust IT connectivity and digitalization for smart cities|Procedure:|We request applicants from the earlier Proposals to align their proposals with the above areas or the Thrust areas given below and resubmit\.|For any query|Terms and Conditions))",
    }

    print(f"--- Text provided for extraction (first 500 chars): {text[:500]} ---") # Display first 500 chars of input for context
    # print(f"DEBUG: Patterns used: {patterns}") # Keep this commented out unless actively debugging regex patterns

    for field, pattern in patterns.items():
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            extracted_value = None
            # Iterate through all capturing groups to find the first non-None value
            for i in range(1, len(match.groups()) + 1):
                if match.group(i) is not None:
                    extracted_value = match.group(i)
                    break

            if extracted_value is not None:
                fields[field] = extracted_value.strip()
                print(f"Extracted '{field}': {fields[field]}")
            else:
                print(f"Could not find '{field}' in the document.")
        else:
            print(f"Could not find '{field}' in the document.")

    return fields


def parse_generated_opportunities(text):
    """
    Parse the Grant Finder AI response into a list of opportunity dicts.

    Preferred format is JSON:
      {
        "opportunities": [
          {"scheme_name": "...", "funding_agency": "...", "last_date_submission": "YYYY-MM-DD", "description": "..."},
          ...
        ]
      }

    Falls back to a best-effort text/regex parser for older/free-form outputs.
    """
    if not text or not isinstance(text, str):
        return []

    # 1) JSON first (most reliable)
    try:
        parsed = json.loads(text)
        if isinstance(parsed, dict) and isinstance(parsed.get("opportunities"), list):
            out = []
            for item in parsed["opportunities"]:
                if not isinstance(item, dict):
                    continue
                out.append(
                    {
                        "scheme_name": str(item.get("scheme_name") or item.get("Programme/Scheme Name") or "N/A").strip(),
                        "funding_agency": str(item.get("funding_agency") or item.get("Funding Agency") or "N/A").strip(),
                        "last_date_submission": str(item.get("last_date_submission") or item.get("Last Date of Submission") or "N/A").strip(),
                        "description": str(item.get("description") or item.get("Description") or "N/A").strip(),
                        "full_text_content": json.dumps(item, ensure_ascii=False),
                    }
                )
            return out
    except Exception:
        pass

    # 2) Text fallback
    # Split into items by numbered list starts (e.g., "1.", "2.") or double newlines.
    chunks = []
    numbered_starts = [m.start() for m in re.finditer(r"(?m)^\s*\d+\.\s+", text)]
    if numbered_starts:
        numbered_starts.append(len(text))
        for i in range(len(numbered_starts) - 1):
            chunk = text[numbered_starts[i] : numbered_starts[i + 1]].strip()
            if chunk:
                chunks.append(chunk)
    else:
        chunks = [c.strip() for c in text.split("\n\n") if c.strip()]

    def _field(pattern, blob):
        m = re.search(pattern, blob, flags=re.IGNORECASE | re.MULTILINE)
        return m.group(1).strip() if m else None

    out = []
    for chunk in chunks:
        scheme = _field(r"Program(?:me)?/Scheme Name\s*[:\-]\s*(.+)$", chunk)
        agency = _field(r"Funding Agency\s*[:\-]\s*(.+)$", chunk)
        deadline = _field(r"Last Date of Submission\s*[:\-]\s*(.+)$", chunk)
        desc = None
        mdesc = re.search(r"(?is)Description\s*[:\-]\s*(.+)$", chunk)
        if mdesc:
            desc = mdesc.group(1).strip()

        out.append(
            {
                "scheme_name": scheme or "N/A",
                "funding_agency": agency or "N/A",
                "last_date_submission": deadline or "N/A",
                "description": desc or "N/A",
                "full_text_content": chunk,
            }
        )
    return out


def fetch_dst_announcements(url=DST_ANNOUNCEMENTS_URL, timeout=20):
    """
    Fetch DST "What's New -> Announcement" page and extract announcement titles + links.
    Returns a list of opportunity dicts compatible with the Grant Finder display.
    """
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) GrantFinder/1.0"
    }
    resp = requests.get(url, headers=headers, timeout=timeout)
    resp.raise_for_status()

    soup = BeautifulSoup(resp.text, "html.parser")

    items = []
    seen = set()

    # Collect all meaningful links (program rows + recent announcements are both link-heavy)
    for a in soup.select("a[href]"):
        title = a.get_text(" ", strip=True)
        href = a.get("href")
        if not title or not href:
            continue
        abs_url = urljoin(url, href)

        # Skip obvious nav/social links; keep PDFs and announcement/program links
        if any(x in abs_url.lower() for x in ["facebook.com", "twitter.com", "youtube.com", "sitemap", "contact", "feedback"]):
            continue
        if len(title) < 8:
            continue

        key = (title, abs_url)
        if key in seen:
            continue
        seen.add(key)

        items.append(
            {
                "scheme_name": title,
                "funding_agency": "Department of Science & Technology (DST)",
                "last_date_submission": "N/A",
                "description": f"Source: {abs_url}",
                "source_url": abs_url,
                "full_text_content": f"{title}\n{abs_url}",
            }
        )

    return items


def fetch_anrf_homepage(url=ANRF_HOMEPAGE_URL, timeout=20):
    """
    Fetch ANRF homepage and extract program/call links.
    Note: The site is interactive; this is a best-effort HTML parse that still yields useful links.
    Returns a list of opportunity dicts compatible with the Grant Finder display.
    """
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) GrantFinder/1.0"
    }
    resp = requests.get(url, headers=headers, timeout=timeout)
    resp.raise_for_status()

    soup = BeautifulSoup(resp.text, "html.parser")

    items = []
    seen = set()

    # Prefer obvious “calls”/“proposal”/program links; fall back to general internal links.
    for a in soup.select("a[href]"):
        title = a.get_text(" ", strip=True)
        href = a.get("href")
        if not title or not href:
            continue
        abs_url = urljoin(url, href)

        # Keep only ANRF internal links
        if "anrfonline.in" not in abs_url.lower():
            continue
        if len(title) < 6:
            continue

        title_l = title.lower()
        keep = any(
            k in title_l
            for k in [
                "call",
                "proposal",
                "grant",
                "fellowship",
                "program",
                "scheme",
                "mission",
                "fund",
            ]
        )
        if not keep:
            continue

        key = (title, abs_url)
        if key in seen:
            continue
        seen.add(key)

        items.append(
            {
                "scheme_name": title,
                "funding_agency": "Anusandhan National Research Foundation (ANRF)",
                "last_date_submission": "N/A",
                "description": f"Source: {abs_url}",
                "source_url": abs_url,
                "full_text_content": f"{title}\n{abs_url}",
            }
        )

    # If parsing yields nothing, still provide the homepage as a single “opportunity” entry.
    if not items:
        items.append(
            {
                "scheme_name": "ANRF Portal - Programs / Calls",
                "funding_agency": "Anusandhan National Research Foundation (ANRF)",
                "last_date_submission": "N/A",
                "description": f"Source: {url}",
                "source_url": url,
                "full_text_content": url,
            }
        )

    return items


def _google_sheet_csv_export_url(sheet_id: str, gid: str) -> str:
    return f"https://docs.google.com/spreadsheets/d/{sheet_id}/export?format=csv&gid={gid}"


def fetch_google_sheet_opportunities(
    sheet_id: str = OPPORTUNITY_SHEET_ID,
    gid: str = OPPORTUNITY_SHEET_GID,
    timeout: int = 30,
    max_bytes: int = 5_000_000,
    max_rows: int = 2000,
):
    """
    Fetch a Google Sheet via CSV export (requires the sheet to be publicly accessible).
    Returns opportunity dicts.
    """
    url = _google_sheet_csv_export_url(sheet_id, gid)
    headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) GrantFinder/1.0"}
    resp = requests.get(url, headers=headers, timeout=timeout, stream=True)
    # If the sheet is not public, Google often returns 403/404 or an HTML interstitial.
    if resp.status_code != 200:
        raise RuntimeError(f"Google Sheet CSV export not accessible (HTTP {resp.status_code}). Make sure the sheet is shared publicly or published to web.")

    buf = bytearray()
    for chunk in resp.iter_content(chunk_size=65536):
        if not chunk:
            continue
        buf.extend(chunk)
        if len(buf) >= max_bytes:
            break

    text = buf.decode("utf-8", errors="replace")
    # If we got HTML instead of CSV, bail with a helpful message.
    if "<html" in text.lower() and "google" in text.lower():
        raise RuntimeError("Google Sheet returned HTML instead of CSV. Please publish the sheet (File → Share/Publish to web) or make it accessible to anyone with the link.")

    f = io.StringIO(text)
    reader = csv.reader(f)
    rows = []
    for i, row in enumerate(reader):
        if i > max_rows:
            break
        rows.append(row)

    if not rows or len(rows) < 2:
        return []

    headers_row = [h.strip() for h in rows[0]]

    def find_col(candidates):
        for cand in candidates:
            for idx, h in enumerate(headers_row):
                if h.lower() == cand.lower():
                    return idx
        # fuzzy contains
        for cand in candidates:
            for idx, h in enumerate(headers_row):
                if cand.lower() in h.lower():
                    return idx
        return None

    name_i = find_col(["Programme/Scheme Name", "Program", "Programme", "Scheme", "Opportunity", "Call", "Title", "Name"])
    agency_i = find_col(["Funding Agency", "Agency", "Sponsor", "Funder"])
    deadline_i = find_col(["Last Date of Submission", "Submission Deadline", "Deadline", "Due Date", "Closing Date", "End Date"])
    desc_i = find_col(["Description", "Summary", "Details", "Notes", "Keywords", "Scope", "Eligibility"])

    # URL may appear in multiple columns; capture the first cell containing http(s)
    url_header_i = find_col(["URL", "Link", "Website", "Web", "Source"])

    items = []
    for row in rows[1:]:
        if not any((c or "").strip() for c in row):
            continue

        def cell(i):
            if i is None or i >= len(row):
                return ""
            return (row[i] or "").strip()

        scheme_name = cell(name_i) or ""
        funding_agency = cell(agency_i) or "N/A"
        last_date = cell(deadline_i) or "N/A"
        description = cell(desc_i) or "N/A"

        # Find a URL: either in the dedicated url column, or anywhere in the row.
        source_url = cell(url_header_i)
        if not source_url or "http" not in source_url.lower():
            for c in row:
                if c and "http" in str(c).lower():
                    source_url = str(c).strip()
                    break

        if not scheme_name:
            # If we can't find a name, skip
            continue

        items.append(
            {
                "scheme_name": scheme_name,
                "funding_agency": funding_agency or "N/A",
                "last_date_submission": last_date or "N/A",
                "description": description or "N/A",
                "source_url": source_url or OPPORTUNITY_SHEET_URL,
                "full_text_content": " | ".join([c.strip() for c in row if (c or "").strip()])[:4000],
            }
        )

    return items


def fetch_india_science_technology_latest(url: str = INDIA_SCI_TECH_LATEST_URL, timeout: int = 25):
    """
    Fetch IndiaScienceAndTechnology 'Latest Updates' and extract likely opportunity/call links.
    Returns a list of opportunity dicts compatible with the Grant Finder display.
    """
    headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) GrantFinder/1.0"}
    try:
        resp = requests.get(url, headers=headers, timeout=timeout)
        resp.raise_for_status()
    except requests.exceptions.SSLError:
        # Some Windows environments lack the CA chain for this site; retry without verification.
        resp = requests.get(url, headers=headers, timeout=timeout, verify=False)
        resp.raise_for_status()

    soup = BeautifulSoup(resp.text, "html.parser")
    items = []
    seen = set()

    for a in soup.select("a[href]"):
        title = a.get_text(" ", strip=True)
        href = a.get("href")
        if not title or not href:
            continue
        abs_url = urljoin(url, href)
        if "indiascienceandtechnology.gov.in" not in abs_url.lower():
            continue
        if len(title) < 8:
            continue

        t = title.lower()
        # Keep items that look like opportunities/calls/announcements; avoid generic navigation
        keep = any(k in t for k in ["call", "proposal", "grant", "fund", "fellow", "fellowship", "invited", "invitation", "applications", "apply", "scheme", "program", "programme", "announcement", "opportunity"])
        if not keep:
            continue

        key = (title, abs_url)
        if key in seen:
            continue
        seen.add(key)

        items.append(
            {
                "scheme_name": title,
                "funding_agency": "India Science & Technology (GoI)",
                "last_date_submission": "N/A",
                "description": f"Source: {abs_url}",
                "source_url": abs_url,
                "full_text_content": f"{title}\n{abs_url}",
            }
        )

    # If nothing was found, still provide the page as a single entry for visibility.
    if not items:
        items.append(
            {
                "scheme_name": "India Science & Technology - Latest Updates",
                "funding_agency": "India Science & Technology (GoI)",
                "last_date_submission": "N/A",
                "description": f"Source: {url}",
                "source_url": url,
                "full_text_content": url,
            }
        )

    return items


def rank_opportunities_by_keywords(opportunities, keywords, top_k=7):
    """
    Lightweight local ranking (no AI): scores opportunities by keyword overlap with title/description.
    """
    if not opportunities:
        return []
    if not keywords:
        return opportunities[:top_k]

    # Normalize keywords
    kw_tokens = []
    for k in keywords:
        if not k:
            continue
        kw_tokens.extend(re.findall(r"[a-z0-9]+", str(k).lower()))
    kw_tokens = [t for t in kw_tokens if len(t) > 2]

    def score(opp):
        hay = f"{opp.get('scheme_name','')} {opp.get('description','')}".lower()
        return sum(hay.count(t) for t in kw_tokens)

    ranked = sorted(opportunities, key=score, reverse=True)
    return ranked[:top_k]


def filter_opportunities_by_keywords(opportunities, keywords):
    """
    Remove opportunities with zero keyword overlap (strict relevance filter).
    """
    if not opportunities:
        return []
    if not keywords:
        return opportunities

    kw_tokens = []
    for k in keywords:
        if not k:
            continue
        kw_tokens.extend(re.findall(r"[a-z0-9]+", str(k).lower()))
    kw_tokens = [t for t in kw_tokens if len(t) > 2]
    if not kw_tokens:
        return opportunities

    def score(opp):
        hay = f"{opp.get('scheme_name','')} {opp.get('description','')}".lower()
        return sum(hay.count(t) for t in kw_tokens)

    return [opp for opp in opportunities if score(opp) > 0]


def _extract_deadline_from_text(text: str) -> Optional[str]:
    """
    Best-effort extraction of a submission deadline from arbitrary call text.
    Returns a string representation (as found) or None.
    """
    if not text:
        return None

    # Prefer dates near typical deadline keywords
    keyword_patterns = [
        r"last\s*date\s*(?:of\s*)?(?:submission|submitting|to\s*submit)",
        r"submission\s*deadline",
        r"last\s*date",
        r"deadline",
        r"due\s*date",
        r"closing\s*date",
    ]

    # Common date formats:
    # - 2026-01-16, 16-01-2026, 16/01/2026
    # - 16 Jan 2026, 16 January 2026
    # - January 16, 2026
    date_patterns = [
        r"(\d{4}-\d{2}-\d{2})",
        r"(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})",
        r"(\d{1,2}\s+(?:Jan|January|Feb|February|Mar|March|Apr|April|May|Jun|June|Jul|July|Aug|August|Sep|Sept|September|Oct|October|Nov|November|Dec|December)\s+\d{2,4})",
        r"((?:Jan|January|Feb|February|Mar|March|Apr|April|May|Jun|June|Jul|July|Aug|August|Sep|Sept|September|Oct|October|Nov|November|Dec|December)\s+\d{1,2},\s*\d{4})",
    ]

    # Scan line-by-line to keep context tight
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    for idx, ln in enumerate(lines):
        ln_l = ln.lower()
        if any(re.search(kp, ln_l) for kp in keyword_patterns):
            window = " ".join(lines[idx : idx + 4])  # include a few lines after
            for dp in date_patterns:
                m = re.search(dp, window, flags=re.IGNORECASE)
                if m:
                    return m.group(1).strip()

    # Fallback: first date anywhere in the text
    blob = "\n".join(lines[:500])  # cap work
    for dp in date_patterns:
        m = re.search(dp, blob, flags=re.IGNORECASE)
        if m:
            return m.group(1).strip()
    return None


def _fetch_text_from_url(url: str, timeout: int = 25) -> str:
    """
    Fetch URL and return extracted text. Supports HTML and PDF (best-effort).
    """
    headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) GrantFinder/1.0"}
    resp = requests.get(url, headers=headers, timeout=timeout)
    resp.raise_for_status()

    content_type = (resp.headers.get("content-type") or "").lower()
    is_pdf = "pdf" in content_type or url.lower().endswith(".pdf")

    if is_pdf:
        try:
            reader = PyPDF2.PdfReader(io.BytesIO(resp.content))
            chunks = []
            for page in reader.pages[:10]:  # cap pages for speed
                try:
                    chunks.append(page.extract_text() or "")
                except Exception:
                    continue
            return "\n".join(chunks)
        except Exception:
            return ""

    # HTML / other: use trafilatura if available, otherwise BeautifulSoup text
    try:
        extracted = trafilatura.extract(resp.text, include_comments=False, include_tables=False)
        if extracted:
            return extracted
    except Exception:
        pass

    soup = BeautifulSoup(resp.text, "html.parser")
    return soup.get_text("\n", strip=True)


def enrich_opportunities_with_deadlines_only(opportunities, max_to_check: int = 10):
    """
    Mutates opportunities in-place: attempts to populate 'last_date_submission'
    from the linked call page/PDF when possible. Uses Streamlit session cache.
    """
    if not opportunities:
        return opportunities

    cache = st.session_state.get("deadline_cache", {})
    if not isinstance(cache, dict):
        cache = {}

    checked = 0
    for opp in opportunities:
        if checked >= max_to_check:
            break
        url = opp.get("source_url")

        if not url:
            continue

        # Pull from caches when possible
        if url in cache and (not opp.get("last_date_submission") or opp.get("last_date_submission") == "N/A"):
            opp["last_date_submission"] = cache[url] or "N/A"
        if opp.get("last_date_submission") and opp.get("last_date_submission") != "N/A":
            continue

        checked += 1
        try:
            txt = _fetch_text_from_url(url)
            deadline = _extract_deadline_from_text(txt) if (not opp.get("last_date_submission") or opp.get("last_date_submission") == "N/A") else None
        except Exception:
            deadline = None

        if deadline is not None:
            cache[url] = deadline
            opp["last_date_submission"] = deadline or "N/A"

    st.session_state["deadline_cache"] = cache
    return opportunities


def _parse_deadline_to_date(deadline_str: str) -> Optional[datetime]:
    """
    Convert a best-effort extracted deadline string into a datetime (date portion used).
    Returns None if parsing fails.
    """
    if not deadline_str or deadline_str == "N/A":
        return None

    s = str(deadline_str).strip()

    # ISO: YYYY-MM-DD
    try:
        return datetime.strptime(s, "%Y-%m-%d")
    except Exception:
        pass

    # Numeric dates: DD-MM-YYYY, DD/MM/YYYY, DD.MM.YYYY, plus 2-digit years.
    # Note: Many sources (esp. US-based) use MM/DD/YYYY, so we apply heuristics by separator.
    m = re.search(r"(\d{1,2})([./-])(\d{1,2})[./-](\d{2,4})", s)
    if m:
        a, sep, b, y = int(m.group(1)), m.group(2), int(m.group(3)), int(m.group(4))
        if y < 100:
            y += 2000
        # Build candidate interpretations
        candidates = []
        # Dot-separated dates are commonly DD.MM.YYYY
        if sep == ".":
            candidates.append((a, b))  # day, month
        # Slash-separated often MM/DD/YYYY in spreadsheets
        elif sep == "/":
            candidates.append((b, a))  # day=second, month=first
            candidates.append((a, b))  # fallback DD/MM
        else:
            # Hyphen: prefer DD-MM, fallback MM-DD
            candidates.append((a, b))
            candidates.append((b, a))

        for day, month in candidates:
            try:
                return datetime(y, month, day)
            except Exception:
                continue

    # "16 Jan 2026" / "16 January 2026"
    m = re.search(
        r"(\d{1,2})\s+"
        r"(Jan|January|Feb|February|Mar|March|Apr|April|May|Jun|June|Jul|July|Aug|August|Sep|Sept|September|Oct|October|Nov|November|Dec|December)"
        r"\s+(\d{2,4})",
        s,
        flags=re.IGNORECASE,
    )
    if m:
        d = int(m.group(1))
        mon = m.group(2).lower()
        y = int(m.group(3))
        if y < 100:
            y += 2000
        months = {
            "jan": 1,
            "january": 1,
            "feb": 2,
            "february": 2,
            "mar": 3,
            "march": 3,
            "apr": 4,
            "april": 4,
            "may": 5,
            "jun": 6,
            "june": 6,
            "jul": 7,
            "july": 7,
            "aug": 8,
            "august": 8,
            "sep": 9,
            "sept": 9,
            "september": 9,
            "oct": 10,
            "october": 10,
            "nov": 11,
            "november": 11,
            "dec": 12,
            "december": 12,
        }
        try:
            return datetime(y, months[mon], d)
        except Exception:
            pass

    # "January 16, 2026"
    m = re.search(
        r"(Jan|January|Feb|February|Mar|March|Apr|April|May|Jun|June|Jul|July|Aug|August|Sep|Sept|September|Oct|October|Nov|November|Dec|December)"
        r"\s+(\d{1,2}),\s*(\d{4})",
        s,
        flags=re.IGNORECASE,
    )
    if m:
        mon = m.group(1).lower()
        d = int(m.group(2))
        y = int(m.group(3))
        months = {
            "jan": 1,
            "january": 1,
            "feb": 2,
            "february": 2,
            "mar": 3,
            "march": 3,
            "apr": 4,
            "april": 4,
            "may": 5,
            "jun": 6,
            "june": 6,
            "jul": 7,
            "july": 7,
            "aug": 8,
            "august": 8,
            "sep": 9,
            "sept": 9,
            "september": 9,
            "oct": 10,
            "october": 10,
            "nov": 11,
            "november": 11,
            "dec": 12,
            "december": 12,
        }
        try:
            return datetime(y, months[mon], d)
        except Exception:
            pass

    return None


def filter_active_open_calls(opportunities, include_no_deadline: bool = True):
    """
    Keep only Active/Open calls:
    - Exclude items explicitly marked CLOSED
    - Exclude items with a parsed deadline that is today or in the past (must be > today)
    - Optionally include items with unknown/no deadline (N/A)
    """
    if not opportunities:
        return []

    today = datetime.now().date()
    filtered = []
    for opp in opportunities:
        title = (opp.get("scheme_name") or "").strip()
        title_l = title.lower()

        if "closed" in title_l:
            continue

        deadline_str = (opp.get("last_date_submission") or "").strip()
        dt = _parse_deadline_to_date(deadline_str)
        if dt:
            if dt.date() <= today:
                continue
            # Store normalized info for display/debug
            opp["deadline_date_iso"] = dt.date().isoformat()
        else:
            if not include_no_deadline:
                continue

        filtered.append(opp)

    return filtered


def extract_alignment_score(report_text):
    match = re.search(r"Alignment Score: (\d+\.?\d*)/10", report_text)
    if match:
        try:
            return float(match.group(1))
        except ValueError:
            return None
    return None

def split_proposal_into_sections(full_proposal_draft, template_sections_str):
    sections = {}
    template_section_titles = [line.strip() for line in template_sections_str.split('\n') if line.strip()]

    # Use a regex to find sections. This assumes section titles are typically at the start of a line and followed by content.
    # It also tries to handle numbered or bulleted sections.
    # This is a heuristic and might need adjustment based on typical proposal formats.
    
    
    # Pattern to match section titles (e.g., 1. Introduction, 2. Background, 2.1. Sub-section)
    # It looks for lines starting with a number and a dot, or a bold markdown header.
    section_title_pattern = re.compile(r"^\s*(\d+\.)+\s*([^\n]+)|^(#+\s*[^\n]+)", re.MULTILINE)

    # Find all potential section starts in the draft
    matches = list(section_title_pattern.finditer(full_proposal_draft))

    for i, match in enumerate(matches):
        # Extract the full matched title string
        if match.group(2): # For (X.X.X. Title) pattern
            current_section_title = match.group(2).strip()
        elif match.group(3): # For (# Title) pattern
            current_section_title = match.group(3).strip().lstrip('# ').strip()
        else:
            continue # Should not happen with the given pattern

        start_pos = match.end()
        end_pos = None

        if i + 1 < len(matches):
            # The end of the current section is the start of the next section
            end_pos = matches[i+1].start()
        else:
            # If it's the last section, it goes to the end of the draft
            end_pos = len(full_proposal_draft)

        section_content = full_proposal_draft[start_pos:end_pos].strip()
        sections[current_section_title] = section_content

    # If no sections were found by regex, or if template sections are provided, 
    # try to align based on template section titles if they are clearly present.
    # This is a fallback/enhancement. For now, rely on regex from draft if template sections are not explicitly markers in the draft.
    # This part can be made more sophisticated if template_sections_str needs to override or guide the split.

    # As a simpler initial implementation, if template_sections_str is available and 
    # a direct regex split of the proposal draft isn't perfect, we can try to find and assign content based on template sections.
    # For now, we'll return what the regex found, and rely on template_sections_str for iteration in brainstorm_room.

    return sections

# --- OpenAI is initialized above with openai_client ---

# --- Database Initialization (Unified) ---
def init_dbs():
    conn = sqlite3.connect(DATABASE_FILE)
    c = conn.cursor()
    # Proposals table (from Grant Proposal Overview Generator)
    c.execute('''
        CREATE TABLE IF NOT EXISTS proposals (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            timestamp TEXT NOT NULL,
            funding_agency TEXT,
            scheme_type TEXT,
            duration TEXT,
            budget TEXT,
            thrust_areas TEXT,
            eligibility TEXT,
            submission_format TEXT,
            user_research_background TEXT,
            template_sections TEXT,
            full_proposal_content TEXT,
            brainstorm_analysis_report TEXT,
            alignment_score FLOAT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        );
    ''')

    # Add brainstorm_analysis_report column if it doesn't exist
    try:
        c.execute('''
            ALTER TABLE proposals ADD COLUMN brainstorm_analysis_report TEXT;
        ''')
    except sqlite3.OperationalError as e:
        if "duplicate column name" not in str(e):
            raise

    # Add created_at column if it doesn't exist
    try:
        c.execute('''
            ALTER TABLE proposals ADD COLUMN created_at TIMESTAMP;
        ''')
        c.execute('''
            UPDATE proposals SET created_at = CURRENT_TIMESTAMP WHERE created_at IS NULL;
        ''')
    except sqlite3.OperationalError as e:
        if "duplicate column name" not in str(e):
            raise

    # Add alignment_score column if it doesn't exist
    try:
        c.execute('''
            ALTER TABLE proposals ADD COLUMN alignment_score FLOAT;
        ''')
    except sqlite3.OperationalError as e:
        if "duplicate column name" not in str(e):
            raise

    # Generated Opportunities table (from Grant Finder)
    c.execute('''
        CREATE TABLE IF NOT EXISTS generated_opportunities (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            timestamp TEXT NOT NULL,
            scheme_name TEXT,
            funding_agency TEXT,
            description TEXT,
            last_date_submission TEXT,
            extracted_keywords TEXT,
            full_text_content TEXT
        )
    ''')

    # Add is_processed column to generated_opportunities if it doesn't exist
    try:
        c.execute('''
            ALTER TABLE generated_opportunities ADD COLUMN is_processed INTEGER DEFAULT 0;
        ''')
    except sqlite3.OperationalError as e:
        if "duplicate column name" not in str(e):
            raise

    # User Research Profiles table (new for storing user profiles independently)
    c.execute('''
        CREATE TABLE IF NOT EXISTS user_profiles (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            timestamp TEXT NOT NULL,
            profile_name TEXT UNIQUE,
            research_background TEXT
        )
    ''')

    conn.commit()
    conn.close()

# Call unified DB init
init_dbs()

def save_proposal_to_db(proposal_data):
    with sqlite3.connect(DATABASE_FILE) as conn:
        cursor = conn.cursor()
        cursor.execute("""
            INSERT INTO proposals (timestamp, funding_agency, scheme_type, duration, budget, thrust_areas, eligibility, submission_format, user_research_background, template_sections, full_proposal_content, brainstorm_analysis_report, alignment_score)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            proposal_data['timestamp'],
            proposal_data['funding_agency'],
            proposal_data['scheme_type'],
            proposal_data['duration'],
            proposal_data['budget'],
            proposal_data['thrust_areas'],
            proposal_data['eligibility'],
            proposal_data['submission_format'],
            proposal_data['user_research_background'],
            proposal_data['template_sections'],
            proposal_data['full_proposal_content'],
            proposal_data['brainstorm_analysis_report'],
            proposal_data.get('alignment_score', None)  # Add alignment_score with a default of None
        ))
        conn.commit()

def save_generated_opportunity_to_db(opportunity_data):
    conn = sqlite3.connect(DATABASE_FILE)
    c = conn.cursor()
    c.execute("""
        INSERT INTO generated_opportunities (timestamp, scheme_name, funding_agency, description, last_date_submission, extracted_keywords, full_text_content)
        VALUES (?, ?, ?, ?, ?, ?, ?)
    """, (
        opportunity_data['timestamp'],
        opportunity_data['scheme_name'],
        opportunity_data['funding_agency'],
        opportunity_data['description'],
        opportunity_data['last_date_submission'],
        opportunity_data['extracted_keywords'],
        opportunity_data['full_text_content']
    ))
    conn.commit()
    conn.close()

def load_all_proposals():
    with sqlite3.connect(DATABASE_FILE) as conn:
        conn.row_factory = sqlite3.Row  # This allows us to access columns by name
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM proposals ORDER BY timestamp DESC")
        proposals = cursor.fetchall()
        return [dict(proposal) for proposal in proposals] # Convert to list of dictionaries

def delete_proposal_from_db(proposal_id):
    with sqlite3.connect(DATABASE_FILE) as conn:
        cursor = conn.cursor()
        cursor.execute("DELETE FROM proposals WHERE id = ?", (proposal_id,))
        conn.commit()
    st.success("Proposal deleted successfully!")

def save_user_profile(profile_name, research_background):
    with sqlite3.connect(DATABASE_FILE) as conn:
        cursor = conn.cursor()
        # Check if profile_name already exists to prevent duplicate entries or to update
        cursor.execute("INSERT OR REPLACE INTO user_profiles (timestamp, profile_name, research_background) VALUES (?, ?, ?)",
                       (datetime.now().strftime("%Y-%m-%d %H:%M:%S"), profile_name, research_background))
        conn.commit()
    st.success(f"Research profile '{profile_name}' saved successfully!")

# --- Helper function to parse brainstorm analysis report ---
def parse_brainstorm_report(report_text):
    sections = {
        "Strengths": "",
        "Weaknesses": "",
        "Recommendations": ""
    }
    current_section = None
    lines = report_text.split('\n')

    section_keywords = {
        "**Strengths**": "Strengths",
        "**Weaknesses**": "Weaknesses",
        "**Recommendations**": "Recommendations"
    }

    for line in lines:
        line_stripped = line.strip()
        found_section = False
        for keyword, section_name in section_keywords.items():
            if line_stripped.startswith(keyword):
                current_section = section_name
                sections[current_section] += line_stripped[len(keyword):].strip() + "\n"
                found_section = True
                break
        if not found_section and current_section and line_stripped:
            sections[current_section] += line_stripped + "\n"
            
    # Clean up trailing newlines
    for key in sections:
        sections[key] = sections[key].strip()
            
    return sections

# --- Helper to load user profiles ---
def load_user_profiles():
    with sqlite3.connect(DATABASE_FILE) as conn:
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()
        cursor.execute("SELECT profile_name, research_background FROM user_profiles ORDER BY timestamp DESC")
        profiles = cursor.fetchall()
        return {profile['profile_name']: profile['research_background'] for profile in profiles}

# --- Helper to delete a user profile ---
def delete_user_profile(profile_name):
    with sqlite3.connect(DATABASE_FILE) as conn:
        cursor = conn.cursor()
        cursor.execute("DELETE FROM user_profiles WHERE profile_name = ?", (profile_name,))
        conn.commit()
    st.success(f"Research profile '{profile_name}' deleted successfully!")

@st.cache_resource
def load_taxonomy(taxonomy_path: Path, file_mtime: Optional[float]):
    if taxonomy_path.exists():
        with open(taxonomy_path, "r", encoding="utf-8") as f:
            return json.load(f)
    st.warning("taxonomy.json not found. Using built-in default taxonomy.")
    return DEFAULT_TAXONOMY

def find_taxonomy_path() -> Path:
    candidates = []
    configured_path = os.getenv("TAXONOMY_PATH") or ""
    if not configured_path:
        try:
            configured_path = st.secrets.get("TAXONOMY_PATH", "")
        except Exception:
            configured_path = ""
    if configured_path:
        candidates.append(Path(configured_path))
    for base_dir in (BASE_DIR, Path.cwd()):
        candidates.append(base_dir / "taxonomy.json")
        candidates.append(base_dir.parent / "taxonomy.json")
        candidates.append(base_dir.parent.parent / "taxonomy.json")
    for candidate in candidates:
        if candidate.exists():
            return candidate
    return candidates[0] if candidates else (BASE_DIR / "taxonomy.json")

taxonomy_path = find_taxonomy_path()

taxonomy_mtime = taxonomy_path.stat().st_mtime if taxonomy_path.exists() else None
taxonomy = load_taxonomy(taxonomy_path, taxonomy_mtime)

# --- Page Config ---
st.set_page_config(page_title="", layout="wide")

# Static top-left logo in the sidebar (exact image, not collapsible)
logo_path = BASE_DIR / "assets" / "idea2impact_logo.png"
if logo_path.exists():
    st.sidebar.image(str(logo_path), use_container_width=False, width=200)
else:
    # If logo is missing, allow one-time upload and persist it
    st.sidebar.markdown('<p style="font-size: 1.3em; font-weight: bold;"></p>', unsafe_allow_html=True)
    uploaded_logo = st.sidebar.file_uploader("Upload logo image (PNG/JPG)", type=["png", "jpg", "jpeg"], key="sidebar_logo_uploader")
    if uploaded_logo is not None:
        try:
            logo_path.parent.mkdir(parents=True, exist_ok=True)
            with open(logo_path, "wb") as f:
                f.write(uploaded_logo.getbuffer())
            st.sidebar.success("Logo saved. Reloading...")
            st.rerun()
        except Exception as e:
            st.sidebar.error(f"Failed to save logo: {e}")

# --- Custom CSS for Cursor-style look ---
st.markdown("""
    <style>
    /* General Streamlit overrides for dark theme */
    .stApp {
        background-color: #1a1a1a;
        color: white; /* Changed for better readability */
    }
    .css-1d391kg, .stButton>button {
        background-color: #2b2b2b;
        color: white;
        border-radius: 4px;
        border: 1px solid #3c3c3c;
    }
    .stTextInput>div>div>input,
    .stTextArea>div>div>textarea,
    .stSelectbox>div>div>div>div,
    .stMultiSelect>div>div>div>div {
        background-color: #3c3c3c;
        color: white;
        border: 1px solid #555;
        border-radius: 4px;
    }
    .stCodeBlock {
        background-color: #2d2d2d;
        color: white;
        border-radius: 5px;
        padding: 1em;
    }
    h1, h2, h3, h4, h5, h6 {
        color: white;
    }
    .stMarkdown {
        font-family: 'Fira Code', monospace; /* Changed for Cursor-style aesthetic */
    }

    /* Custom sidebar navigation styling */
    .sidebar .sidebar-content {
        background-color: #1e1e1e;
        padding-top: 2rem;
    }
    .css-1lcbmhc, .css-zbnxdr {
        background-color: #1e1e1e;
    }
    .css-1oe5zfg {
        padding-top: 2rem;
    }
    .css-1ymn5ad {
        padding-bottom: 2rem;
    }
    .css-1v3fvcr {
        font-weight: bold;
        color: white;
    }
    .css-pkbujm {
        color: white; /* Sidebar link color */
        font-size: 1.1em;
    }
    .css-pkbujm:hover {
        color: #61afef; /* Hover color */
    }
    /* Sidebar item active state */
    .css-1y4fgqg.eqr7sfu4 {
        background-color: #2b2b2b;
        color: #61afef;
        border-left: 3px solid #61afef;
    }

    /* Card-like button styling for the main dashboard */
    .card-container {
        display: flex;
        flex-wrap: wrap; /* Allow cards to wrap */
        gap: 20px; /* Space between cards */
        justify-content: space-between; /* Distribute cards and fill space */
        padding: 20px 0;
    }
    .stCard { /* Streamlit's internal card class for the dashboard cards */
        background-color: #252526; /* Darker background for cards */
        border: 1px solid #333;
        border-radius: 4px;
        box-shadow: 0 4px 8px rgba(0, 0, 0, 0.2);
        transition: transform 0.2s, box-shadow 0.2s;
        min-height: 150px; /* Ensure cards have some height */
        display: flex;
        flex-direction: column;
        justify-content: space-between;
        cursor: pointer;
        flex-grow: 1; /* Allow cards to grow and fill available space */
        flex-basis: 48%; /* Roughly two cards per row with gap */
        max-width: 48%; /* Max width to ensure two per row */
    }
    .stCard:hover {
        transform: translateY(-5px);
        box-shadow: 0 8px 16px rgba(0, 0, 0, 0.4);
    }
    .card-title {
        color: white; /* Use brighter white for titles */
        font-size: 1.2em;
        font-weight: bold;
        margin-bottom: 10px;
    }
    .card-description {
        color: white; /* Slightly lighter grey for description */
        font-size: 0.9em;
        flex-grow: 1; /* Allow description to take up available space */
    }
    .card-button {
        background-color: #007acc; /* VS Code blue for button */
        color: white;
        padding: 10px 15px;
        border: none;
        border-radius: 4px;
        cursor: pointer;
        font-size: 0.9em;
        align-self: flex-start; /* Align button to start of flex item */
        margin-top: 15px; /* Space above the button */
    }
    .card-button:hover {
        background-color: #005f99;
    }

    /* General Streamlit button styling (for main content area) */
    .stButton>button {
        background-color: #007acc; /* VS Code blue */
        color: white;
        border-radius: 4px;
        border: none;
        padding: 10px 20px;
        font-size: 1em;
        cursor: pointer;
        transition: background-color 0.2s;
    }
    .stButton>button:hover {
        background-color: #005f99;
    }
    .stButton>button:focus {
        box-shadow: none;
    }

    /* Text input and text area styling */
    .stTextInput>div>div>input, .stTextArea>div>div>textarea {
        background-color: #333;
        color: white !important;
        border: 1px solid #555;
        border-radius: 4px;
        padding: 8px 12px;
    }
    .stTextInput>div>div>input:focus, .stTextArea>div>div>textarea:focus {
        border-color: #007acc;
        box-shadow: 0 0 0 1px #007acc;
    }

    /* Selectbox styling */
    .stSelectbox>div>div>div {
        background-color: #333;
        color: white;
        border: 1px solid #555;
        border-radius: 4px;
    }
    .stSelectbox>div>div>div>div>span {
        color: white;
    }

    /* Multiselect styling */
    .stMultiSelect>div>div>div {
        background-color: #333;
        color: white;
        border: 1px solid #555;
        border-radius: 4px;
    }
    .stMultiSelect span {
        color: white;
    }
    .stMultiSelect .st-emotion-cache-1bzx45r { /* Chips */
        background-color: #007acc;
        color: white;
    }
    
    /* Expander styling */
    .st-emotion-cache-lq6x5h { /* Target the expander header */
        background-color: #252526;
        border: 1px solid #333;
        border-radius: 4px;
        padding: 10px;
        margin-bottom: 10px;
        color: white;
    }
    .st-emotion-cache-lq6x5h .st-emotion-cache-pkj102 { /* Expander icon */
        color: white;
    }
    /* Custom Card Styles for Dashboard */
    .card-container {
        display: flex;
        flex-wrap: wrap; /* Allow cards to wrap */
        gap: 20px; /* Space between cards */
        justify-content: space-around; /* Distribute cards and fill space */
        padding: 20px 0;
    }
    .stCard { /* Streamlit's internal card class */
        background-color: #252526; /* Darker background for cards */
        border: 1px solid #333;
        border-radius: 4px;
        box-shadow: 0 4px 8px rgba(0, 0, 0, 0.2);
        transition: transform 0.2s, box-shadow 0.2s;
        min-height: 150px; /* Ensure cards have some height */
        display: flex;
        flex-direction: column;
        justify-content: space-between;
        cursor: pointer;
        width: calc(50% - 10px); /* Two cards per row, accounting for gap */
        /* Removed max-width to allow expansion */
    }
    .stCard:hover {
        transform: translateY(-5px);
        box-shadow: 0 8px 16px rgba(0, 0, 0, 0.4);
    }
    .card-title {
        color: white;
        font-size: 1.2em;
        font-weight: bold;
        margin-bottom: 10px;
    }
    .card-description {
        color: white;
        font-size: 0.9em;
    }
    .card-button {
        background-color: #007acc; /* VS Code blue for button */
        color: white;
        padding: 10px 15px;
        border: none;
        border-radius: 4px;
        cursor: pointer;
        font-size: 0.9em;
        align-self: flex-start; /* Align button to start of flex item */
    }
    .card-style {
        background-color: #34285b;
        color: white;
        border-radius: 0.75rem;
        padding: 1.5rem;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.3);
        transition: transform 0.2s;
        cursor: pointer;
        /* text-align: center; will be inline */
    }
    .card-style:hover {
        transform: scale(1.02);
    }
    .card-style .card-title-text {
        font-size: 1.25rem;
        font-weight: 600;
        margin-bottom: 0.5rem;
    }
    .card-style .card-description-text {
        font-size: 0.875rem;
        color: white;
    }
    /* Ensure all labels are white */
    .st-emotion-cache-nahz7x, .st-emotion-cache-vk3357,
    div[data-testid="stWidgetLabel"] label, /* Target labels directly */
    div[data-testid="stWidgetLabel"] p, /* Target text within labels for text_area, etc. */
    .stRadio > label, /* For st.radio main label */
    .stRadio div[role="radio"] > div > p, /* For individual radio button labels */
    .stRadio div[data-testid="stOption"] > div > p, /* For radio button options themselves */
    .stRadio div[data-testid="stOption"] > div > span, /* For radio button options text */
    .stTextArea label, /* Explicitly target text area labels */
    .stTextArea div[data-testid="stWidgetLabel"] p, /* More specific for text area labels */
    .stSelectbox > label, /* For st.selectbox main label */
    .stSelectbox span, /* For st.selectbox selected value and option text */
    .stSelectbox div[role="option"] > div > span /* For individual selectbox options text */
    {
        color: white !important;
    }
    .stRadio > label > div > div > p {
        color: #FFFFFF !important; /* White color */
    }
    div[data-testid="stRadio"] label span {
        color: #FFFFFF !important; /* White color */
    }
    </style>
""", unsafe_allow_html=True)

# --- Session State Initialization ---
if 'current_main_view' not in st.session_state:
    st.session_state['current_main_view'] = 'dashboard' # Default to Dashboard as active
if 'proposal_inputs' not in st.session_state:
    st.session_state['proposal_inputs'] = {}
if 'generated_opportunities' not in st.session_state:
    st.session_state['generated_opportunities'] = []
if 'last_loaded_opportunity' not in st.session_state:
    st.session_state['last_loaded_opportunity'] = None
if 'uploaded_template_sections' not in st.session_state:
    st.session_state['uploaded_template_sections'] = []
if 'user_research_profile' not in st.session_state:
    st.session_state['user_research_profile'] = ""
if 'funding_agency' not in st.session_state:
    st.session_state['funding_agency'] = ""
if 'scheme_type' not in st.session_state:
    st.session_state['scheme_type'] = ""
if 'thrust_areas' not in st.session_state:
    st.session_state['thrust_areas'] = ""
if 'eligibility' not in st.session_state:
    st.session_state['eligibility'] = ""
if 'full_proposal_draft' not in st.session_state:
    st.session_state['full_proposal_draft'] = ""
if 'actual_template_sections_used' not in st.session_state:
    st.session_state['actual_template_sections_used'] = ""
if 'template_upload_error' not in st.session_state:
    st.session_state['template_upload_error'] = False

# Initialize databases on startup
init_dbs()

# --- Navigation Functions ---
def nav_to(view_name):
    st.session_state['current_main_view'] = view_name
    st.rerun()

# --- Sidebar Navigation ---
st.sidebar.markdown('<p style="font-size: 1.5em; font-weight: bold; color: #f0f0f0;"> </p>', unsafe_allow_html=True)
sidebar_items = [
    {"label": "Idea2Impact Studio", "icon": "", "view": "dashboard"},
    {"label": "Grant Finder", "icon": "", "view": "grant_finder"},
    {"label": "Align with My Research", "icon": "", "view": "align_research"},
    {"label": "Grant Proposal Overview", "icon": "", "view": "proposal_generator"},  
    {"label": "Brainstorm Room", "icon": "", "view": "brainstorm_room"},
    {"label": "Draft Final Proposal", "icon": "✍️", "view": "draft_final"},
    {"label": "My Drafts & Submissions", "icon": "", "view": "my_drafts"},
    {"label": "Export & Share Center", "icon": "", "view": "export_share"}
]

for item in sidebar_items:
    # Use st.sidebar.button and apply custom CSS via markdown for styling
    if st.sidebar.button(
        f"{item['icon']} {item['label']}",
        key=f"sidebar_nav_{item['view']}",
        help=f"Go to {item['label']}"
    ):
        nav_to(item['view'])
    
    # Apply active style using markdown for the button's parent div
    if st.session_state['current_main_view'] == item['view']:
        st.sidebar.markdown(f"""
            <style>
                div[data-testid="stSidebarNav"] button[key="sidebar_nav_{item['view']}"] {{
                    background-color: #2b2b2b;
                    color: #61afef;
                    border-left: 3px solid #61afef;
                    font-weight: bold;
                }}
            </style>
        """, unsafe_allow_html=True)

# --- Main Content Area ---
st.header("Idea2Impact Studio")

if st.session_state['current_main_view'] == 'dashboard':
    st.markdown("<h3>Welcome to your AI-powered Grant Writing Assistant!</h3>", unsafe_allow_html=True)

    # Define dashboard cards with titles, icons, and descriptions
    dashboard_cards = [
        {"title": " Explore  Opportunities ", "desc": "Launch the Grant Proposal Overview Generator", "emoji": "", "view": "proposal_generator"},
        {"title": "Align With Expertise", "desc": "Browse active funding calls", "emoji": "", "view": "grant_finder"},
        {"title": " Brainstorm Room ", "desc": "Collaborative brainstorming on proposal ideas", "emoji": "", "view": "brainstorm_room"},
        {"title": "✍️ Draft Final ", "desc": "Generate improved proposal from feedback", "emoji": "", "view": "draft_final"},
        {"title": "Draft Proposal", "desc": "View and manage your saved drafts", "emoji": "", "view": "my_drafts"},
    ]

    # Apply global CSS for overall app styling and custom button styling
    st.markdown('''
    <style>
    .stApp {
        background-color: #121212; /* Apply to the main Streamlit app container */
    }

    /* Container for the card and hidden button */
    .card-wrapper {
        position: relative;
        width: 100%;
        height: 14rem; /* Fixed height for consistency */
        margin-bottom: 1rem; /* Space between cards when they stack */
    }

    /* Styling for Streamlit buttons to make them look like cards, but with transparent background */
    .stButton>button {
        background-color: transparent; /* Make the actual button transparent */
        color: transparent; /* Make button text transparent */
        border: none;
        border-radius: 0.75rem;
        padding: 0;
        box-shadow: none;
        transition: none;
        cursor: pointer;
        width: 100%;
        height: 100%;
        position: absolute; /* Position over the entire wrapper */
        top: 0;
        left: 0;
        z-index: 2; /* Ensure it's above the styled div */
    }
    /* Hide the default hover effect of the Streamlit button */
    .stButton>button:hover {
        transform: none;
        background-color: transparent; 
        color: transparent;
        box-shadow: none;
    }

    /* Styling for the visible card background and content */
    .visible-card {
        background-color: #34285b;
        color: white;
        border-radius: 0.75rem;
        padding: 1.5rem;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.3);
        transition: transform 0.2s;
        cursor: pointer;
        text-align: center;
        height: 100%; /* Fill the wrapper height */
        display: flex;
        flex-direction: column;
        justify-content: center;
        position: relative; /* For z-index to work */
        z-index: 1; /* Ensure it's below the transparent button */
    }
    .visible-card:hover {
        transform: scale(1.02);
    }

    .card-title-text {
        font-size: 1.25rem;
        font-weight: 600;
        margin-bottom: 0.5rem;
        text-align: center; /* Ensure title is centered */
    }
    .card-description-text {
        font-size: 0.875rem;
        color: #a0a0a0;
        text-align: center; /* Ensure description is centered */
    }
    </style>
    ''', unsafe_allow_html=True)

    # Render cards using Streamlit's column system with styled buttons
    cols = st.columns(len(dashboard_cards)) # Create columns based on the number of cards

    for i, card in enumerate(dashboard_cards):
        with cols[i]: # Place each card in its corresponding column
            st.markdown(f"""
                <div class="card-wrapper">
                    <div class="visible-card">
                        <h3 class="card-title-text">{card['emoji']} {card['title']}</h3>
                        <p class="card-description-text">{card['desc']}</p>
                    </div>
                </div>
            """, unsafe_allow_html=True)

            # Create a transparent Streamlit button on top of the visible card
            st.button("", key=f"dashboard_card_{card['view']}", on_click=nav_to, args=(card['view'],))

    st.markdown("---")
    st.markdown("### Getting Started Guide")
    st.code("""
# Welcome to Idea2Impact Studio!

# Use this space to:
# 1. Discover funding opportunities tailored to your research.
# 2. Analyze your Research Profile alignment with funding calls.
# 3. Generate new grant proposals with AI assistance.
# 4. Export professional PDF reports.

# To begin, select an option from the sidebar.
# Happy Grant Writing!
    """, language='python')

elif st.session_state['current_main_view'] == 'proposal_generator':
    # Initialize proposal_generated flag
    if 'proposal_generated' not in st.session_state:
        st.session_state.proposal_generated = False

    st.title("📄 Grant Proposal Overview Generator")

    # --- Section 1: Funding Call Analysis ---
    st.subheader("1. Funding Call Analysis")

    funding_call_option = st.radio(
        "How would you like to provide the funding call details?",
        ("Upload PDF", "Paste Text", "Enter URL", "Select from Saved Opportunities", "Use Analyzed Call (Alignment)")
    )

    # Initialize alignment analysis session state if not present
    if 'alignment_analysis' not in st.session_state:
        st.session_state['alignment_analysis'] = {
            'strategic_recommendations': '',
            'keywords_themes': ''
        }
    
    # Initialize user_research_profile for proposal generator
    if 'user_research_profile' not in st.session_state:
        st.session_state.user_research_profile = ""

    funding_call_text = ""
    # Define a callback function to clear the text input
    def clear_add_section_input():
        st.session_state['tmpl_add_new'] = ""

    if funding_call_option == "Upload PDF":
        uploaded_file = st.file_uploader("Upload PDF of Funding Call", type=["pdf"])
        if uploaded_file:
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text()
            funding_call_text = text
            st.success("PDF uploaded and extracted successfully!")
    elif funding_call_option == "Paste Text":
        funding_call_text = st.text_area("Paste Funding Call Text Here", height=300)
    elif funding_call_option == "Enter URL":
        url = st.text_input("Enter URL of Funding Call")
        if url:
            try:
                # Use trafilatura for better text extraction
                downloaded = trafilatura.fetch_url(url)
                if downloaded:
                    text = trafilatura.extract(downloaded, favor_recall=True)
                    funding_call_text = text if text else "Could not extract sufficient text from the URL."
                else:
                    st.warning("Could not fetch content from the provided URL.")

            except Exception as e:
                st.error(f"Error fetching or parsing URL: {e}")
                funding_call_text = ""
            if funding_call_text:
                st.success("Content extracted from URL!")
    elif funding_call_option == "Select from Saved Opportunities":
        conn = sqlite3.connect(DATABASE_FILE)
        c = conn.cursor()
        c.execute("SELECT id, scheme_name, funding_agency, description, last_date_submission FROM generated_opportunities ORDER BY timestamp DESC")
        opportunities_from_db = c.fetchall()
        conn.close()

        opportunity_options = [("Select an opportunity", None)] + [(f"{o[1]} - {o[2]}", o) for o in opportunities_from_db]
        selected_opportunity_tuple = st.selectbox("Choose a generated funding opportunity", opportunity_options, format_func=lambda x: x[0], key="analysis_opportunity_selector")

        if selected_opportunity_tuple[1] is not None:
            o_data = selected_opportunity_tuple[1]
            selected_opportunity_data = {
                "scheme_name": o_data[1],
                "funding_agency": o_data[2],
                "description": o_data[3],
                "last_date_submission": o_data[4] if len(o_data) > 4 else "N/A"
            }
            # Explicitly update session state variables
            st.session_state['funding_agency'] = selected_opportunity_data.get('funding_agency', '')
            st.session_state['scheme_type'] = selected_opportunity_data.get('scheme_name', '')
            st.session_state['thrust_areas'] = selected_opportunity_data.get('description', '')
            # Assuming eligibility and submission_format are not directly in generated opportunities
            st.session_state['eligibility'] = st.session_state.get('eligibility', 'N/A') 
            st.session_state['submission_format'] = st.session_state.get('submission_format', 'N/A')
            
            funding_call_text = f"""Funding Agency: {selected_opportunity_data.get('funding_agency', '')}
            Scheme Type: {selected_opportunity_data.get('scheme_name', '')}
            Thrust Areas: {selected_opportunity_data.get('description', '')}
            Last Date of Submission: {selected_opportunity_data.get('last_date_submission', '')}"""
            st.success("Loaded funding call details from saved opportunities.")

    elif funding_call_option == "Use Analyzed Call (Alignment)":
        align_data = st.session_state.get('align_selected_opportunity')
        alignment_analysis_data = st.session_state.get('alignment_analysis')

        if align_data:
            st.session_state['funding_agency'] = align_data.get('funding_agency', '')
            st.session_state['scheme_type'] = align_data.get('scheme_name', '')
            st.session_state['thrust_areas'] = align_data.get('description', '')
            st.session_state['last_date_submission'] = align_data.get('last_date_submission', '')
            # Ensure eligibility and submission format are carried over or set to N/A if not available
            st.session_state['eligibility'] = align_data.get('eligibility', 'N/A') # Assuming it might be part of align_data
            st.session_state['submission_format'] = align_data.get('submission_format', 'N/A') # Assuming it might be part of align_data

            # Construct funding_call_text for fields extraction, if needed
            funding_call_text = f"""Funding Agency: {st.session_state['funding_agency']}
            Scheme Type: {st.session_state['scheme_type']}
            Thrust Areas: {st.session_state['thrust_areas']}
            Last Date of Submission: {st.session_state['last_date_submission']}"""

            st.success("Loaded funding call details from Call Alignment.")

            if alignment_analysis_data:
                # Store or update in session state for editing and drafting
                st.session_state['strategic_recommendations'] = alignment_analysis_data.get('strategic_recommendations', '')
                st.session_state['keywords_themes'] = alignment_analysis_data.get('keywords_themes', '')

        else:
            st.info("No analyzed call found. Go to 'Align with My Research' and generate an alignment to populate this option.")

    if funding_call_text:
        st.subheader("Extracted Funding Call Details:")

        # Display extracted info (and allow user to edit)
        st.write(f"DEBUG: funding_call_text content: {funding_call_text[:500]}...") # New Debug line
        st.write(f"DEBUG: Input to extract_fields: {funding_call_text[:1000]}...") # Debugging line
        extracted_fields = extract_fields(funding_call_text)
        st.write(f"DEBUG: Extracted fields result: {extracted_fields}") # Debugging line

        # Convert extracted fields to session state
        st.session_state['scheme_type'] = extracted_fields.get('Scheme Type', '')

        # Explicitly update session state for all extracted fields
        st.session_state['funding_agency'] = extracted_fields.get('Funding Agency', '')
        st.session_state['duration'] = extracted_fields.get('Duration', '')
        st.session_state['budget'] = extracted_fields.get('Budget', '')
        st.session_state['thrust_areas'] = extracted_fields.get('Thrust Areas', '')
        st.session_state['eligibility'] = extracted_fields.get('Eligibility', '')
        st.session_state['submission_format'] = extracted_fields.get('Submission Format', '')
        st.session_state['last_date_submission'] = extracted_fields.get('Last Date of Submission', '')
        st.session_state['scheme_call_name'] = extracted_fields.get('Scheme or Call Name', '')
        st.session_state['scope_objective'] = extracted_fields.get('Scope or Objective of the Programme', '')

        st.markdown(f"**Funding Agency:** {st.session_state['funding_agency']}")
        st.markdown(f"**Scheme or Call Name:** {st.session_state['scheme_call_name']}")
        st.markdown(f"**Scheme Type:** {st.session_state['scheme_type']}")
        st.markdown(f"**Duration:** {st.session_state['duration']}")
        st.markdown(f"**Budget:** {st.session_state['budget']}")
        st.markdown(f"**Thrust Areas:** {st.session_state['thrust_areas']}")
        st.markdown(f"**Scope or Objective of the Programme:** {st.session_state['scope_objective']}")
        st.markdown(f"**Eligibility Criteria:** {st.session_state['eligibility']}")
        st.markdown(f"**Submission Format:** {st.session_state['submission_format']}")
        st.markdown(f"**Last Date of Submission:** {st.session_state['last_date_submission']}")

        # Display editable alignment recommendations and keywords if available
        if funding_call_option == "Use Analyzed Call (Alignment)" and st.session_state.get('alignment_analysis'):
            st.markdown("---")
            st.subheader("Alignment Analysis Insights (Editable):")
            st.session_state['strategic_recommendations'] = st.text_area(
                "Strategic Recommendations",
                value=st.session_state.get('strategic_recommendations', ''),
                height=200,
                key="prop_gen_strategic_recs"
            )
            st.session_state['keywords_themes'] = st.text_area(
                "Keywords / Themes",
                value=st.session_state.get('keywords_themes', ''),
                height=100,
                key="prop_gen_keywords_themes"
            )

        if st.button('Load Last Generated Opportunity from Research Opportunities Generator'):
            conn = sqlite3.connect(DATABASE_FILE)
            c = conn.cursor()
            c.execute("SELECT id, scheme_name, funding_agency, last_date_submission, description FROM generated_opportunities WHERE is_processed = 0 ORDER BY timestamp DESC LIMIT 1")
            last_opportunity = c.fetchone()
            conn.close()

            if last_opportunity:
                opp_id, scheme_name, funding_agency_opp, last_date_submission, description = last_opportunity
                st.session_state['funding_agency'] = funding_agency_opp
                st.session_state['scheme_type'] = scheme_name
                st.session_state['duration'] = f"Deadline: {last_date_submission}"
                st.session_state['thrust_areas'] = description
                st.session_state['submission_format'] = f"Refer to funding agency website by {last_date_submission}"
                st.session_state['eligibility'] = st.session_state.get('eligibility', 'N/A') # Also ensure eligibility is set
                st.success(f"Loaded opportunity: '{scheme_name}' from Research Opportunities Generator. Remember to review details for accuracy!")
                
                conn = sqlite3.connect(DATABASE_FILE)
                c = conn.cursor()
                c.execute("UPDATE generated_opportunities SET is_processed = 1 WHERE id = ?", (opp_id,))
                conn.commit()
                conn.close()
                st.rerun() # Rerun to update the input fields with loaded data
            else:
                st.info("No new generated opportunities to load.")

    user_research_background = st.text_area("Describe your research background and key expertise (Publications, Patents, Grants, etc.)", 
                                         value=st.session_state.get('user_research_profile', ''), 
                                         height=200, 
                                         key="prop_gen_user_research_background")
    st.session_state['user_research_profile'] = user_research_background

    # --- Section 2: Proposal Template & Drafting ---
    st.subheader("2. Proposal Template & Drafting") # Changed from 3 to 2

    template_option = st.radio(
        "How would you like to define your proposal template?",
        ("Generate from Funding Call", "Provide Custom Template Sections", "Upload Funding Agency Template")
    )

    template_sections_input = ""
    if template_option == "Generate from Funding Call":
        if funding_call_text:
            if st.button("Generate Template Sections from Call", key="gen_template_from_call"):
                template_prompt = f"""
                Analyze the following funding call text and extract the typical sections required for a research proposal submission. List them in a clear, numbered or bulleted format. Focus on major sections like 'Introduction', 'Objectives', 'Methodology', 'Budget', 'Timeline', 'Expected Outcomes', 'Bibliography', etc.

                Funding Call Text:
                {funding_call_text}
                """
                with st.spinner("Generating template sections..."):
                    template_response = generate_content_with_retry(SELECTED_MODEL, template_prompt)
                    if template_response:
                        st.session_state['template_sections_generated'] = template_response.text
                        st.session_state['final_template_sections'] = template_response.text # Update final sections too
                        st.success("Template sections generated!")
                    else:
                        st.error("Failed to generate template sections. Please try again.")
            if 'template_sections_generated' in st.session_state:
                st.write(st.session_state['template_sections_generated'])
                template_sections_input = st.session_state['template_sections_generated']
        else:
            st.info("Please provide funding call details first to generate a template.")
    elif template_option == "Provide Custom Template Sections":
        template_sections_input = st.text_area(
            "Enter your custom proposal sections (one per line)",
            height=200,
            value=st.session_state.get('template_sections_generated', '')
        )
    elif template_option == "Upload Funding Agency Template":
        st.info("Upload the official proposal template from the funding agency. Accepted formats: DOCX, PDF, TXT.")
        uploaded_template = st.file_uploader("Upload Template from Funding Call", type=["docx", "pdf", "txt"], key="template_uploader")
        
        # Clear previously parsed sections if a new file is uploaded or option changes
        if uploaded_template != st.session_state.get('last_uploaded_template_file'):
            st.session_state['uploaded_template_sections_parsed'] = []
            st.session_state['last_uploaded_template_file'] = uploaded_template
            st.session_state['template_upload_error'] = False # Reset error state
            if uploaded_template is not None: # If a new file is uploaded, clear any previous template_sections_generated
                st.session_state['template_sections_generated'] = ""

        if uploaded_template is not None:
            file_type = os.path.splitext(uploaded_template.name)[1].lower()
            extracted_template_text = ""
            sections_list = []

            try:
                if file_type == ".txt":
                    extracted_template_text = uploaded_template.read().decode(errors='ignore')
                elif file_type == ".pdf":
                    pdf_reader = PyPDF2.PdfReader(uploaded_template)
                    text_accum = ""
                    for page in pdf_reader.pages:
                        text_accum += page.extract_text() or "\n"
                    extracted_template_text = text_accum
                elif file_type == ".docx":
                    if docx is None:
                        st.error("DOCX support requires python-docx. Please retry after it installs or upload PDF/TXT.")
                        st.session_state['template_upload_error'] = True
                        raise ValueError("docx not installed")
                    else:
                        doc = docx.Document(uploaded_template)
                        extracted_template_text = "\n".join(p.text for p in doc.paragraphs)
                else:
                    st.error("Unsupported template format. Please upload a .docx, .pdf, or .txt file.")
                    st.session_state['template_upload_error'] = True
                    raise ValueError("Unsupported file type")

                if extracted_template_text:
                    # Use AI to extract structured sections from the uploaded template content
                    with st.spinner("Analyzing uploaded template for sections..."):
                        template_analysis_prompt = f"""
                        Analyze the following document text and extract the main section headings and their hierarchical structure. List them clearly. Focus on major sections (e.g., Introduction, Methodology, Budget, Conclusion).

                        Document Text:
                        {extracted_template_text}

                        Provide the output as a numbered list of section titles. Example:
                        1. Introduction
                        2. Background and Significance
                            2.1. Preliminary Studies
                            2.2. Gaps in Literature
                        3. Specific Aims
                        4. Research Design and Methods
                            4.1. Aim 1 Methods
                            4.2. Aim 2 Methods
                        5. Budget
                        6. Timeline
                        7. Expected Outcomes
                        8. References
                        """
                        analysis_response = generate_content_with_retry(SELECTED_MODEL, template_analysis_prompt)
                        if analysis_response and analysis_response.text:
                            sections_list = [line.strip() for line in analysis_response.text.split('\n') if line.strip()]
                            st.session_state['uploaded_template_sections_parsed'] = sections_list
                            st.success("Template sections extracted and ready for use!")
                        else:
                            st.warning("Could not automatically extract sections from the uploaded template using AI. Please use 'Provide Custom Template Sections' instead.")
                            st.session_state['template_upload_error'] = True
                elif not extracted_template_text:
                    st.warning("No text could be extracted from the uploaded file. Please ensure it contains readable text.")
                    st.session_state['template_upload_error'] = True

            except Exception as e:
                if not st.session_state['template_upload_error']:
                    st.error(f"Failed to process uploaded template: {e}")
                st.session_state['template_upload_error'] = True

        # Display parsed sections if available and no error
        if st.session_state.get('uploaded_template_sections_parsed') and not st.session_state['template_upload_error']:
            st.markdown("### Parsed Template Sections (Read-only):")
            for section in st.session_state['uploaded_template_sections_parsed']:
                st.markdown(f"- {section}")
            template_sections_input = "\n".join(st.session_state['uploaded_template_sections_parsed'])
            st.session_state['final_template_sections'] = template_sections_input # Set final sections
        elif st.session_state['template_upload_error']:
            st.warning("Displaying default editable sections due to template upload/parsing issues.")
            # Fallback to custom sections input if there was an error with upload/parsing
            template_sections_input = st.text_area(
                "Enter your custom proposal sections (one per line)",
                height=200,
                value=st.session_state.get('template_sections_generated', '')
            )
        else:
            st.info("Upload a template to see its parsed sections here.")

    # Only show editable sections for 'Generate from Funding Call' or 'Provide Custom Template Sections'
    # or if there was an error with 'Upload Funding Agency Template'
    if st.session_state.proposal_generated and (template_option != "Upload Funding Agency Template" or st.session_state.get('template_upload_error')):
        st.markdown("### Template Sections (editable)")
        sections_for_edit = st.session_state['uploaded_template_sections']

        # Render each section with edit/delete and move controls
        indices_to_delete = []
        for idx, section in enumerate(sections_for_edit):
            cols = st.columns([6, 1, 1, 1])
            with cols[0]:
                new_val = st.text_input(f"Section {idx+1}", value=section, key=f"tmpl_sec_{idx}")
                sections_for_edit[idx] = new_val
            with cols[1]:
                if st.button("↑", key=f"move_up_{idx}") and idx > 0:
                    sections_for_edit[idx-1], sections_for_edit[idx] = sections_for_edit[idx], sections_for_edit[idx-1]
                    st.session_state['uploaded_template_sections'] = sections_for_edit
                    st.rerun()
            with cols[2]:
                if st.button("↓", key=f"move_down_{idx}") and idx < len(sections_for_edit)-1:
                    sections_for_edit[idx+1], sections_for_edit[idx] = sections_for_edit[idx], sections_for_edit[idx+1]
                    st.session_state['uploaded_template_sections'] = sections_for_edit
                    st.rerun()
            with cols[3]:
                if st.button("Delete", key=f"delete_{idx}"):
                    indices_to_delete.append(idx)
        # Apply deletions after loop
        if indices_to_delete:
            st.session_state['uploaded_template_sections'] = [s for i, s in enumerate(sections_for_edit) if i not in indices_to_delete]

        # Add new section control
        new_section_text = st.text_input("Add new section", value=st.session_state.get('tmpl_add_new', ''), key="tmpl_add_new")
        if st.button("Add Section", key="tmpl_add_btn", on_click=clear_add_section_input):
            if new_section_text.strip():
                st.session_state['uploaded_template_sections'].append(new_section_text.strip())
                st.rerun()
            else:
                st.warning("Section title cannot be empty.")

        # Confirm sections to use
        if st.button("Use These Sections", key="tmpl_use_sections"):
            template_sections_input = "\n".join(st.session_state['uploaded_template_sections'])
            st.session_state['template_sections_generated'] = template_sections_input
            st.session_state['final_template_sections'] = template_sections_input
            st.success("Template sections set for drafting.")

    # Only update final_template_sections if we have a non-empty value to avoid clobbering on rerender
    if template_sections_input:
        st.session_state['final_template_sections'] = template_sections_input

    if st.button("Proposal overview Draft"):
        if st.session_state.get('final_template_sections') and st.session_state.get('user_research_profile'):
            # Include alignment recommendations and keywords if they exist in session state
            alignment_insights_for_prompt = ""
            if st.session_state.get('strategic_recommendations'):
                alignment_insights_for_prompt += f"\n\nStrategic Recommendations: {st.session_state['strategic_recommendations']}"
            if st.session_state.get('keywords_themes'):
                alignment_insights_for_prompt += f"\n\nKeywords/Themes to Emphasize: {st.session_state['keywords_themes']}"

            # Determine the template sections to use for the prompt
            template_sections_to_use = st.session_state.get('final_template_sections', '')
            if st.session_state.get('uploaded_template_sections_parsed') and not st.session_state.get('template_upload_error'):
                template_sections_to_use = "\n".join(st.session_state['uploaded_template_sections_parsed'])
            
            st.write(f"DEBUG (Proposal Generator): template_sections_to_use length: {len(template_sections_to_use)}")

            # Also store the actually used template sections in session state for saving/exporting
            st.session_state['actual_template_sections_used'] = template_sections_to_use

            proposal_draft_prompt = f"""
            Generate a comprehensive research proposal draft based on the following information. Structure the proposal according to the provided template sections. Incorporate insights from the alignment report and the user's research background.

            ---
            Funding Call Details:
            Funding Agency: {st.session_state.get('funding_agency', 'N/A')}
            Scheme Type: {st.session_state.get('scheme_type', 'N/A')}
            Duration: {st.session_state.get('duration', 'N/A')}
            Budget (suggested if available): {st.session_state.get('budget', 'N/A')}
            Thrust Areas: {st.session_state.get('thrust_areas', 'N/A')}
            Eligibility: {st.session_state.get('eligibility', 'N/A')}
            Submission Format: {st.session_state.get('submission_format', 'N/A')}

            ---
            User Research Background:
            {st.session_state.get('user_research_profile', '')}

            ---
            Alignment Report & Ideas:
            # Removed direct use of alignment_report, as insights are now structured.
            {alignment_insights_for_prompt}

            ---
            Proposal Template Sections (write content for each):
            {template_sections_to_use}

            ---
            Instructions for AI:
            - Write a detailed and persuasive proposal.
            - Ensure logical flow and coherence between sections.
            - Use academic and professional language.
            - Highlight novelty, feasibility, and potential impact.
            - For sections like 'Budget' or 'Timeline', provide realistic placeholders or general statements if specific figures are not derivable from the input, or suggest what should be included.
            - If a section like 'Bibliography' is listed, just put a placeholder like "[References/Bibliography to be added]"
            - Ensure the content directly addresses the funding call's requirements and aligns with the user's background.
            - The full proposal should be at least 1500 words, but ideally around 2000-3000 words for a substantial draft.
            """
            with st.spinner("Generating full proposal draft (this may take a few minutes for a comprehensive draft)..."):
                full_proposal_response = generate_content_with_retry(SELECTED_MODEL, proposal_draft_prompt)
                if full_proposal_response:
                    st.session_state['full_proposal_draft'] = full_proposal_response.text
                    st.success("Full proposal draft generated!")
                    st.session_state.proposal_generated = True # Set flag to True
                    st.write(f"DEBUG (Proposal Generator): full_proposal_draft length: {len(st.session_state.get('full_proposal_draft', ''))}")
                    st.write(f"DEBUG (Proposal Generator): actual_template_sections_used populated: {bool(st.session_state.get('actual_template_sections_used'))}")
                else:
                    st.error("Failed to generate full proposal draft. Please try again.")
        else:
            st.warning("Please ensure funding call details, research background, and template sections are provided.")

    if st.session_state.proposal_generated and 'full_proposal_draft' in st.session_state:
        st.markdown("### Full Proposal Draft:")
        st.write(st.session_state['full_proposal_draft'])

        # --- Section 4: Export & Save ---
        st.subheader("4. Export & Save")

        if st.button("Save Proposal to Database"):
            proposal_data = {
                'timestamp': datetime.now().isoformat(),
                'funding_agency': st.session_state.get('funding_agency', ''),
                'scheme_type': st.session_state.get('scheme_type', ''),
                'duration': st.session_state.get('duration', ''),
                'budget': st.session_state.get('budget', ''),
                'thrust_areas': st.session_state.get('thrust_areas', ''),
                'eligibility': st.session_state.get('eligibility', ''),
                'submission_format': st.session_state.get('submission_format', ''),
                'user_research_background': st.session_state.get('user_research_profile', ''),
                'template_sections': st.session_state.get('actual_template_sections_used', ''), # Use the actually used sections
                'full_proposal_content': st.session_state.get('full_proposal_draft', ''),
                'brainstorm_analysis_report': st.session_state.get('brainstorm_analysis_report', ''),
                'alignment_score': st.session_state.get('alignment_score', None)  # Add alignment_score with a default of None
            }
            save_proposal_to_db(proposal_data)
            st.success("Proposal saved to database!")

        # Export as PDF
        if st.button("Export Proposal as PDF"):
            pdf = PDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)

            # Write template sections as headings and then the content
            sections_content = st.session_state.get('full_proposal_draft', '')
            sections_list = st.session_state.get('actual_template_sections_used', '').split('\n')
            
            # Simple heuristic to try and match content to sections
            # This can be made more sophisticated with AI parsing of the draft
            current_content_idx = 0
            for section_title in sections_list:
                if section_title.strip():
                    pdf.set_font("Arial", 'B', 14)
                    pdf.multi_cell(0, 10, section_title)
                    pdf.set_font("Arial", size=12)
                    pdf.ln(2) # Small gap after section title
                    
                    # Find content for this section (very basic approach)
                    # In a real scenario, the AI should generate content structured by sections
                    # For now, we'll just flow the entire draft content
                    pdf.multi_cell(0, 10, sections_content) # For now, dump all content
                    break # Only dump full content once for simplicity

            pdf_output = pdf.output(dest='S').encode('latin-1')
            st.download_button(
                label="Download Proposal PDF",
                data=pdf_output,
                file_name="generated_proposal.pdf",
                mime="application/pdf"
            )

        # --- Retrieve & Load Proposals ---
        st.subheader("5. Retrieve & Load Saved Proposals")
        
        saved_proposals = load_all_proposals()
        
        if saved_proposals:
            # Create a user-friendly display for selection
            proposal_options = []
            for idx, p in enumerate(saved_proposals):
                title = f"{p['timestamp']} - {p['funding_agency']} - {p['scheme_type']}"
                proposal_options.append((title, idx)) # Store (display_text, original_index)

            # Create a mapping from display text to original index
            option_to_index = {title: idx for title, idx in proposal_options}
            display_titles = [title for title, _ in proposal_options]

            selected_title = st.selectbox("Select a saved proposal to load:", display_titles)

            if selected_title:
                original_idx = option_to_index[selected_title]
                selected_proposal_data = saved_proposals[original_idx]
                
                st.markdown(f"**Selected Proposal Details:**")
                st.json({
                    "Funding Agency": selected_proposal_data.get('funding_agency', 'N/A'),
                    "Scheme Type": selected_proposal_data.get('scheme_type', 'N/A'),
                    "Timestamp": selected_proposal_data.get('timestamp', 'N/A')
                })

                col_load, col_delete = st.columns([1, 1])
                with col_load:
                    if st.button("Load Selected Proposal Content"):
                        # Update session state with loaded proposal data
                        st.session_state['funding_agency'] = selected_proposal_data.get('funding_agency', '')
                        st.session_state['scheme_type'] = selected_proposal_data.get('scheme_type', '')
                        st.session_state['duration'] = selected_proposal_data.get('duration', '')
                        st.session_state['budget'] = selected_proposal_data.get('budget', '')
                        st.session_state['thrust_areas'] = selected_proposal_data.get('thrust_areas', '')
                        st.session_state['eligibility'] = selected_proposal_data.get('eligibility', '')
                        st.session_state['submission_format'] = selected_proposal_data.get('submission_format', '')
                        st.session_state['user_research_profile'] = selected_proposal_data.get('user_research_background', '')
                        st.session_state['actual_template_sections_used'] = selected_proposal_data.get('template_sections', '')
                        st.session_state['full_proposal_draft'] = selected_proposal_data.get('full_proposal_content', '')
                        st.session_state['brainstorm_analysis_report'] = selected_proposal_data.get('brainstorm_analysis_report', '')
                        
                        st.success(f"Proposal '{selected_title}' loaded successfully into the generator!")
                        st.rerun() # Rerun to update the UI with loaded data
                with col_delete:
                    if st.button("Delete Selected Proposal"):
                        proposal_id_to_delete = selected_proposal_data.get('id')
                        if proposal_id_to_delete:
                            delete_proposal_from_db(proposal_id_to_delete)
                            st.rerun()
                        else:
                            st.error("Could not retrieve proposal ID for deletion.")
        else:
            st.info("No proposals saved yet. Save a proposal above to see it here!")

elif st.session_state['current_main_view'] == 'grant_finder':
    st.header("🔍 Grant Finder (AI-Powered Research Opportunity Generator)")
    
    st.markdown("### Discover New Funding Opportunities with AI")
    st.write("Select research areas and let AI find relevant opportunities for you.")

    active_only = st.checkbox(
        "Only show Active/Open calls (exclude past-deadline calls)",
        value=True,
        help="Shows only calls with a parsed deadline strictly after today (and not marked CLOSED).",
    )

    # Use st.columns to place broad and specific domain selection side-by-side
    col_broad, col_specific = st.columns(2)

    with col_broad:
        selected_broad_domain = st.selectbox("Select Broad Domain", list(taxonomy.keys()))

    selected_specific_areas = []
    if selected_broad_domain:
        with col_specific:
            selected_specific_areas = st.multiselect(
                "Select Specific Research Areas",
                taxonomy[selected_broad_domain]
            )
    else:
        st.info("Please select a broad domain first.")

    if st.button("Generate Research Opportunities"):
        if selected_specific_areas:
            combined_areas = ", ".join(selected_specific_areas)
            
            # Calculate minimum submission date (today + 15 days)
            min_submission_date = (datetime.now() + timedelta(days=15)).strftime("%Y-%m-%d")

            st.info("Searching opportunities... This may take a moment.")
            with st.spinner("Fetching opportunities from DST announcements, ANRF portal, Opportunity Sheet, and IndiaScienceAndTechnology..."):
                dst_items, anrf_items, sheet_items, ist_items = [], [], [], []
                dst_err, anrf_err, sheet_err, ist_err = None, None, None, None
                try:
                    dst_items = fetch_dst_announcements()
                except Exception as e:
                    dst_err = e
                try:
                    anrf_items = fetch_anrf_homepage()
                except Exception as e:
                    anrf_err = e
                try:
                    sheet_items = fetch_google_sheet_opportunities()
                except Exception as e:
                    sheet_err = e
                try:
                    ist_items = fetch_india_science_technology_latest()
                except Exception as e:
                    ist_err = e

            merged = (dst_items or []) + (anrf_items or []) + (sheet_items or []) + (ist_items or [])
            if merged:
                st.caption(
                    f"Sources fetched — DST: {len(dst_items or [])}, ANRF: {len(anrf_items or [])}, "
                    f"Sheet: {len(sheet_items or [])}, IndiaS&T: {len(ist_items or [])}."
                )
                ranked = rank_opportunities_by_keywords(merged, selected_specific_areas, top_k=50)
                ranked = filter_opportunities_by_keywords(ranked, selected_specific_areas)
                # Try to read submission deadlines from the call links (best-effort)
                with st.spinner("Reading Last Date of Submission from call links (best-effort)..."):
                    ranked = enrich_opportunities_with_deadlines_only(ranked, max_to_check=30)
                if active_only:
                    # Debug counts to explain why zero results can happen
                    total = len(ranked)
                    parsed_deadlines = sum(1 for r in ranked if _parse_deadline_to_date((r.get("last_date_submission") or "").strip()))
                    unknown_deadlines = sum(1 for r in ranked if not _parse_deadline_to_date((r.get("last_date_submission") or "").strip()))
                    filtered = filter_active_open_calls(ranked, include_no_deadline=False)
                    st.caption(
                        f"Active/Open filter: total {total}, parsed deadlines {parsed_deadlines}, unknown deadlines {unknown_deadlines}, active {len(filtered)}."
                    )
                    ranked = filtered
                    if not ranked:
                        st.warning(
                            "No active opportunities matched your filters. "
                            "This usually means all parsed deadlines are before today or missing."
                        )
                st.session_state['generated_opportunities_raw'] = json.dumps(
                    {
                        "sources": [
                            DST_ANNOUNCEMENTS_URL,
                            ANRF_HOMEPAGE_URL,
                            OPPORTUNITY_SHEET_URL,
                            INDIA_SCI_TECH_LATEST_URL,
                        ],
                        "opportunities": ranked,
                    },
                    ensure_ascii=False,
                    indent=2,
                )
                st.session_state['generated_opportunities'] = ranked
                st.success("Opportunities generated from Diffrent sources!")
                if dst_err:
                    st.warning(f"DST fetch issue (partial results still shown): {dst_err}")
                if anrf_err:
                    st.warning(f"ANRF fetch issue (partial results still shown): {anrf_err}")
                if sheet_err:
                    st.warning(
                        "Opportunity Sheet fetch issue (partial results still shown). "
                        f"To enable it, make the sheet public or publish it to web. Error: {sheet_err}"
                    )
                if ist_err:
                    st.warning(f"IndiaScienceAndTechnology fetch issue (partial results still shown): {ist_err}")
            else:
                # Fallback: keep the app usable even if DST fetch fails (e.g., offline)
                st.warning("DST/ANRF opportunities unavailable right now. Falling back to AI-generated ideas (may be fictional).")
                prompt = f"""Generate a list of 5-7 innovative and actionable research opportunities or project ideas based on the following research areas: {combined_areas}.

Return ONLY valid JSON (no markdown, no code fences) with this exact shape:
{{
  "opportunities": [
    {{
      "scheme_name": "string",
      "funding_agency": "string",
      "last_date_submission": "YYYY-MM-DD (prefer {min_submission_date} or later if you invent a date, otherwise 'N/A')",
      "description": "2-3 sentences"
    }}
  ]
}}
"""
                with st.spinner("AI is brainstorming opportunities..."):
                    response = generate_content_with_retry(SELECTED_MODEL, prompt)
                if response and response.text:
                    raw_text = response.text.strip()
                    st.session_state['generated_opportunities_raw'] = raw_text
                    st.session_state['generated_opportunities'] = parse_generated_opportunities(raw_text)
                    if active_only:
                        st.session_state['generated_opportunities'] = filter_active_open_calls(
                            st.session_state['generated_opportunities'], include_no_deadline=True
                        )
                    st.success("Opportunities generated!")
                else:
                    st.error("Failed to generate research opportunities. Please try again.")
        else:
            st.warning("Please select at least one specific research area.")

    if 'generated_opportunities' in st.session_state and st.session_state['generated_opportunities']:
        st.subheader("Generated Research Opportunities:")
        opps = st.session_state.get('generated_opportunities')
        # Backwards compatibility: if older runs stored a string, parse it now.
        if isinstance(opps, str):
            st.session_state['generated_opportunities_raw'] = opps
            opps = parse_generated_opportunities(opps)
            st.session_state['generated_opportunities'] = opps

        # If parsing failed, show raw output to make debugging easy.
        if not opps:
            st.warning("Couldn't parse the AI response into structured opportunities. Showing raw output:")
            raw_text = st.session_state.get('generated_opportunities_raw', '')
            st.code(raw_text or "(empty response)")
        else:
            for i, opp in enumerate(opps):
                st.markdown(f"**Opportunity {i+1}**")

                opportunity_data = {
                    "timestamp": datetime.now().isoformat(),
                    "scheme_name": opp.get("scheme_name", "N/A"),
                    "funding_agency": opp.get("funding_agency", "N/A"),
                    "last_date_submission": opp.get("last_date_submission", "N/A"),
                    "description": opp.get("description", "N/A"),
                    "extracted_keywords": st.session_state.get('grant_finder_keywords', ''),
                    "full_text_content": opp.get("full_text_content", "")
                }

                st.write(f"**Programme/Scheme Name:** {opportunity_data['scheme_name']}")
                st.write(f"**Funding Agency:** {opportunity_data['funding_agency']}")
                st.write(f"**Last Date of Submission:** {opportunity_data['last_date_submission']}")
                if opp.get("source_url"):
                    st.markdown(f"**Source Link:** {opp.get('source_url')}")

                if st.button(f"Submit this Opportunity ({i+1}) to Grant Proposal Overview Generator", key=f"submit_opp_{i}"):
                    st.info(f"Attempting to submit opportunity: {opportunity_data['scheme_name']}")
                    save_generated_opportunity_to_db(opportunity_data)
                    st.success(f"Opportunity '{opportunity_data['scheme_name']}' submitted! Go to 'Grant Proposal  Overview Generator' and click 'Load Last Generated Opportunity'.")
                    st.rerun()
                st.markdown("---") # Separator
elif st.session_state['current_main_view'] == 'align_research':
    st.header("🎯 Align with My Research")
    st.markdown("### Align Your Research Profile with Funding Opportunities")
    st.write("Provide your research background and select a funding opportunity to see how well they align.")

    st.subheader("1. Your Research Background")
    
    # Initialize session state for user_research_profile if not already present
    if 'user_research_profile' not in st.session_state:
        st.session_state.user_research_profile = ""

    user_research_profile = st.text_area("Describe your research background and key expertise (Publicatons, Patent, Grants,etc)", 
                                         height=250, 
                                         key="align_research_profile_input",
                                         value=st.session_state.user_research_profile)

    # Update session state when the text area changes
    st.session_state.user_research_profile = user_research_profile

    st.subheader("2. Provide or Select a Funding Opportunity")
    
    # Removed custom CSS for radio button text color as global theme is applied
    # st.markdown("""
    # <style>
    # .stRadio > label {
    #     color: #FFFFFF !important; /* White color */
    # }
    # div[data-testid="stRadio"] label span {
    #     color: #FFFFFF !important; /* White color */
    # }
    # </style>
    # """, unsafe_allow_html=True)

    funding_call_option_align = st.radio(
        "How would you like to provide the funding call details for alignment?",
        ("Select from Saved Opportunities", "Paste Text", "Upload PDF", "Enter URL"),
        key="align_funding_call_option"
    )

    funding_call_text_align = ""
    selected_opportunity_data = {} # Initialize as dictionary to store extracted fields

    if funding_call_option_align == "Select from Saved Opportunities":
        conn = sqlite3.connect(DATABASE_FILE)
        c = conn.cursor()
        c.execute("SELECT id, scheme_name, funding_agency, description, last_date_submission FROM generated_opportunities ORDER BY timestamp DESC")
        opportunities_from_db = c.fetchall()
        conn.close()

        opportunity_options = [("Select an opportunity", None)] + [(f"{o[1]} - {o[2]}", o) for o in opportunities_from_db]
        selected_opportunity_tuple = st.selectbox("Choose a generated funding opportunity", opportunity_options, format_func=lambda x: x[0], key="align_opportunity_selector")

        if selected_opportunity_tuple[1] is not None:
            o_data = selected_opportunity_tuple[1]
            selected_opportunity_data = {
                "scheme_name": o_data[1],
                "funding_agency": o_data[2],
                "description": o_data[3],
                "last_date_submission": o_data[4] if len(o_data) > 4 else "N/A" # Assuming last_date_submission is the 5th element (index 4)
            }
            st.markdown("**Selected Opportunity Details (Editable):**")
            edited_scheme_name = st.text_input("Scheme or Call Name", value=selected_opportunity_data['scheme_name'], key="align_scheme_name")
            edited_funding_agency = st.text_input("Funding Agency", value=selected_opportunity_data['funding_agency'], key="align_funding_agency")
            edited_description = st.text_area("Thrust, Scope or Objective of the Programme", value=selected_opportunity_data['description'], height=150, key="align_description")
            edited_last_date = st.text_input("Last Date of Submission", value=selected_opportunity_data['last_date_submission'], key="align_last_date_submission")
            selected_opportunity_data['scheme_name'] = edited_scheme_name
            selected_opportunity_data['funding_agency'] = edited_funding_agency
            selected_opportunity_data['description'] = edited_description
            selected_opportunity_data['last_date_submission'] = edited_last_date

            # Deletion control for the selected generated opportunity
            selected_opportunity_id = o_data[0]
            if st.button("Delete this opportunity", key=f"delete_generated_opportunity_{selected_opportunity_id}"):
                try:
                    conn = sqlite3.connect(DATABASE_FILE)
                    c = conn.cursor()
                    c.execute("DELETE FROM generated_opportunities WHERE id = ?", (selected_opportunity_id,))
                    conn.commit()
                    conn.close()
                    st.success("Selected opportunity deleted.")
                    st.rerun()
                except Exception as e:
                    st.error(f"Failed to delete opportunity: {e}")

    elif funding_call_option_align == "Paste Text":
        funding_call_text_align = st.text_area("Paste Funding Call Text Here", height=300, key="align_paste_text")
        if funding_call_text_align:
            st.subheader("Extracted Funding Call Details:")
            st.write(f"DEBUG: Input to extract_fields (align paste): {funding_call_text_align[:1000]}") # Debugging line
            extracted_fields = extract_fields(funding_call_text_align)
            st.write(f"DEBUG: Extracted fields result (align paste): {extracted_fields}") # Debugging line
            selected_opportunity_data = {
                "scheme_name": extracted_fields.get('Scheme Type', 'N/A'),
                "funding_agency": extracted_fields.get('Funding Agency', 'N/A'),
                "description": extracted_fields.get('Thrust Areas', 'N/A'), # Using Thrust Areas as a general description
                "last_date_submission": extracted_fields.get('Last Date of Submission', 'N/A')
            }
            edited_scheme_name = st.text_input("Scheme or Call Name", value=selected_opportunity_data['scheme_name'], key="align_scheme_name")
            edited_funding_agency = st.text_input("Funding Agency", value=selected_opportunity_data['funding_agency'], key="align_funding_agency")
            edited_description = st.text_area("Thrust, Scope or Objective of the Programme", value=selected_opportunity_data['description'], height=150, key="align_description")
            edited_last_date = st.text_input("Last Date of Submission", value=selected_opportunity_data['last_date_submission'], key="align_last_date_submission")
            selected_opportunity_data['scheme_name'] = edited_scheme_name
            selected_opportunity_data['funding_agency'] = edited_funding_agency
            selected_opportunity_data['description'] = edited_description
            selected_opportunity_data['last_date_submission'] = edited_last_date
            st.success("Text extracted and fields processed!")

            # Update global session state for Brainstorm Room
            st.session_state['funding_agency'] = selected_opportunity_data.get('funding_agency', '')
            st.session_state['scheme_type'] = selected_opportunity_data.get('scheme_name', '')
            st.session_state['thrust_areas'] = selected_opportunity_data.get('description', '')
            st.session_state['eligibility'] = st.session_state.get('eligibility', 'N/A')

    elif funding_call_option_align == "Upload PDF":
        uploaded_file = st.file_uploader("Upload PDF of Funding Call", type=["pdf"], key="align_upload_pdf")
        if uploaded_file:
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text()
            funding_call_text_align = text
            st.success("PDF uploaded and extracted successfully!")
            if funding_call_text_align:
                st.subheader("Extracted Funding Call Details:")
                st.write(f"DEBUG: Input to extract_fields (align PDF): {funding_call_text_align[:1000]}") # Debugging line
                extracted_fields = extract_fields(funding_call_text_align)
                st.write(f"DEBUG: Extracted fields result (align PDF): {extracted_fields}") # Debugging line
                selected_opportunity_data = {
                    "scheme_name": extracted_fields.get('Scheme Type', 'N/A'),
                    "funding_agency": extracted_fields.get('Funding Agency', 'N/A'),
                    "description": extracted_fields.get('Thrust Areas', 'N/A'),
                    "last_date_submission": extracted_fields.get('Last Date of Submission', 'N/A')
                }
                edited_scheme_name = st.text_input("Scheme or Call Name", value=selected_opportunity_data['scheme_name'], key="align_scheme_name")
                edited_funding_agency = st.text_input("Funding Agency", value=selected_opportunity_data['funding_agency'], key="align_funding_agency")
                edited_description = st.text_area("Thrust, Scope or Objective of the Programme", value=selected_opportunity_data['description'], height=150, key="align_description")
                edited_last_date = st.text_input("Last Date of Submission", value=selected_opportunity_data['last_date_submission'], key="align_last_date_submission")
                selected_opportunity_data['scheme_name'] = edited_scheme_name
                selected_opportunity_data['funding_agency'] = edited_funding_agency
                selected_opportunity_data['description'] = edited_description
                selected_opportunity_data['last_date_submission'] = edited_last_date

                # Update global session state for Brainstorm Room
                st.session_state['funding_agency'] = selected_opportunity_data.get('funding_agency', '')
                st.session_state['scheme_type'] = selected_opportunity_data.get('scheme_name', '')
                st.session_state['thrust_areas'] = selected_opportunity_data.get('description', '')
                st.session_state['eligibility'] = st.session_state.get('eligibility', 'N/A')

    elif funding_call_option_align == "Enter URL":
        url = st.text_input("Enter URL of Funding Call", key="align_url_input")
        if url:
            try:
                downloaded = trafilatura.fetch_url(url)
                if downloaded:
                    text = trafilatura.extract(downloaded, favor_recall=True)
                    funding_call_text_align = text if text else "Could not extract sufficient text from the URL."
                else:
                    st.warning("Could not fetch content from the provided URL.")

            except Exception as e:
                st.error(f"Error fetching or parsing URL: {e}")
                funding_call_text_align = ""
            if funding_call_text_align and funding_call_text_align != "Could not extract sufficient text from the URL.":
                st.success("Content extracted from URL!")
                st.subheader("Extracted Funding Call Details:")
                st.write(f"DEBUG: Input to extract_fields (align URL): {funding_call_text_align[:1000]}") # Debugging line
                extracted_fields = extract_fields(funding_call_text_align)
                st.write(f"DEBUG: Extracted fields result (align URL): {extracted_fields}") # Debugging line
                selected_opportunity_data = {
                    "scheme_name": extracted_fields.get('Scheme Type', 'N/A'),
                    "funding_agency": extracted_fields.get('Funding Agency', 'N/A'),
                    "description": extracted_fields.get('Thrust Areas', 'N/A'),
                    "last_date_submission": extracted_fields.get('Last Date of Submission', 'N/A')
                }
                edited_scheme_name = st.text_input("Scheme or Call Name", value=selected_opportunity_data['scheme_name'], key="align_scheme_name")
                edited_funding_agency = st.text_input("Funding Agency", value=selected_opportunity_data['funding_agency'], key="align_funding_agency")
                edited_description = st.text_area("Thrust, Scope or Objective of the Programme", value=selected_opportunity_data['description'], height=150, key="align_description")
                edited_last_date = st.text_input("Last Date of Submission", value=selected_opportunity_data['last_date_submission'], key="align_last_date_submission")
                selected_opportunity_data['scheme_name'] = edited_scheme_name
                selected_opportunity_data['funding_agency'] = edited_funding_agency
                selected_opportunity_data['description'] = edited_description
                selected_opportunity_data['last_date_submission'] = edited_last_date

                # Update global session state for Brainstorm Room
                st.session_state['funding_agency'] = selected_opportunity_data.get('funding_agency', '')
                st.session_state['scheme_type'] = selected_opportunity_data.get('scheme_name', '')
                st.session_state['thrust_areas'] = selected_opportunity_data.get('description', '')
                st.session_state['eligibility'] = st.session_state.get('eligibility', 'N/A')
            elif funding_call_text_align == "Could not extract sufficient text from the URL.":
                st.warning("No significant text could be extracted from the provided URL.")

    # This block ensures that if data is successfully extracted, it is used for alignment
    if st.button("Generate Alignment Analysis", key="generate_alignment_button"):
        if user_research_profile and selected_opportunity_data:
            st.session_state['align_user_profile'] = user_research_profile
            st.session_state['align_selected_opportunity'] = selected_opportunity_data

            alignment_prompt = f"""
            As an expert grant evaluator, critically analyze the alignment between the provided Research Profile and Funding Opportunity.

            Instructions:
            1. Summarize the core expertise and contributions of the research profile (based only on the provided information, assume publications are implicitly part of 'Research Profile' text if not explicitly separated).
            2. Map this expertise to the funding call's stated priorities, explicitly distinguishing between:
               - Direct alignment (clear fit with call objectives)
               - Indirect or speculative alignment (possible applications if reframed)
               - Non-alignment (areas with no overlap)
            3. Highlight major gaps that reduce alignment (e.g., domain mismatch, lack of collaborations, lack of translational/clinical orientation).
            4. Suggest strategies to increase alignment (e.g., reframing expertise, building collaborations, translational roadmaps).
            5. Assign a critical **Alignment Score (0–10)**, where:
               - 0–3 = Very weak/no alignment
               - 4–6 = Moderate alignment (requires strong reframing)
               - 7–8 = Strong alignment (with clear fit and collaborations)
               - 9–10 = Excellent alignment (highly competitive)
            6. Maintain a professional, analytical tone.

            --- Output Structure ---
            Research Profile Summary:
            [Summary of core expertise and contributions]

            Alignment with Call Priorities:
            - Direct Alignment: [Points of direct fit]
            - Indirect/Speculative Alignment: [Possible applications if reframed]
            - Non-Alignment: [Areas with no overlap]

            Key Gaps:
            [Major gaps reducing alignment]

            Strategic Recommendations:
            [Strategies to increase alignment]

            Alignment Score: X.X/10
            ---

            Research Profile:
            {user_research_profile}

            Funding Opportunity:
            Scheme Name: {selected_opportunity_data['scheme_name']}
            Funding Agency: {selected_opportunity_data['funding_agency']}
            Description: {selected_opportunity_data['description']}
            """

            # Truncate prompt to avoid exceeding token limits
            max_prompt_length = 25000  # Adjust as needed based on model limits
            if len(alignment_prompt) > max_prompt_length:
                st.warning(f"Alignment prompt truncated from {len(alignment_prompt)} to {max_prompt_length} characters.")
                alignment_prompt = alignment_prompt[:max_prompt_length]

            st.write(f"DEBUG: Alignment prompt length: {len(alignment_prompt)}") # For debugging

            with st.spinner("Generating alignment analysis report..."):
                alignment_response = generate_content_with_retry(SELECTED_MODEL, alignment_prompt)
                if alignment_response:
                    st.session_state['alignment_analysis_report'] = alignment_response.text
                    st.success("Alignment analysis generated successfully!")
                else:
                    st.error("Failed to generate alignment analysis. Please try again.")
        else:
            st.warning("Please provide your research background and select or provide a funding opportunity.")

    if 'alignment_analysis_report' in st.session_state:
        col_left, col_right = st.columns([4, 1])
        with col_left:
            st.subheader("Alignment Analysis Report")
        with col_right:
            analysis_text = st.session_state['alignment_analysis_report']
            analysis_score = extract_alignment_score(analysis_text)
            # Store alignment score in session state for saving to database
            st.session_state['alignment_score'] = analysis_score
            score_value = f"{analysis_score:.1f}/10" if analysis_score is not None else "N/A"
            st.metric(label="Alignment Score", value=score_value)
        st.write(st.session_state['alignment_analysis_report'])

        # Parse and store strategic recommendations and keywords/themes
        strategic_recs_match = re.search(r'(?is)Strategic\s*Recommendations:\s*(.*?)(?:\n\d+\.?\s*Keywords/Themes:|Alignment Score:|\Z)', analysis_text)
        keywords_themes_match = re.search(r'(?is)Keywords/Themes:\s*(.*?)(?:\n\d+\.?\s*Alignment Score:|\Z)', analysis_text)

        strategic_recs = strategic_recs_match.group(1).strip() if strategic_recs_match else ''
        keywords_themes = keywords_themes_match.group(1).strip() if keywords_themes_match else ''

        st.session_state['alignment_analysis'] = {
            'strategic_recommendations': strategic_recs,
            'keywords_themes': keywords_themes
        }

elif st.session_state['current_main_view'] == 'brainstorm_room':
    st.header("🧠 Brainstorm Room")
    st.markdown("### Get AI-Powered Feedback on Your Proposal Sections")

    # --- Input Section ---
    st.subheader("1. Provide Your Proposal Details for Brainstorming")

    # Option to load saved proposals
    saved_proposals = load_all_proposals()
    if saved_proposals:
        proposal_options = [("Select a saved proposal to load", None)]
        for p in saved_proposals:
            title = f"{p['timestamp']} - {p['funding_agency']} - {p['scheme_type']}"
            proposal_options.append((title, p))
        
        selected_proposal_for_brainstorm = st.selectbox(
            "Load from Saved Proposals", 
            proposal_options, 
            format_func=lambda x: x[0], 
            key="br_load_saved_proposal_selector"
        )

        if selected_proposal_for_brainstorm[1] is not None:
            if st.button("Load Selected Proposal into Brainstorm Workspace", key="br_load_proposal_button"):
                loaded_data = selected_proposal_for_brainstorm[1]
                st.session_state['full_proposal_draft'] = loaded_data.get('full_proposal_content', '')
                st.session_state['actual_template_sections_used'] = loaded_data.get('template_sections', '')
                st.session_state['user_research_profile'] = loaded_data.get('user_research_background', '')
                st.session_state['funding_agency'] = loaded_data.get('funding_agency', '')
                st.session_state['scheme_type'] = loaded_data.get('scheme_type', '')
                st.session_state['thrust_areas'] = loaded_data.get('thrust_areas', '')
                st.session_state['eligibility'] = loaded_data.get('eligibility', '')
                st.success(f"Proposal '{loaded_data.get('scheme_type', 'N/A')}' loaded successfully!")
                st.rerun() # Rerun to update the displayed content
    else:
        st.info("No saved proposals found. Generate and save a proposal in 'Grant Proposal Overview Generator' first.")

    st.markdown("---")
    
    # Separate input sections for funding call and proposal
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("2. Funding Call Details")
        funding_call_option_br = st.radio(
            "How would you like to provide funding call details?",
            ("Paste Text", "Upload PDF", "Enter URL"),
            key="br_funding_call_option"
        )
        
        funding_call_text_br = ""
        if funding_call_option_br == "Paste Text":
            funding_call_text_br = st.text_area("Paste Funding Call Text Here", height=200, key="br_paste_funding_call")
        elif funding_call_option_br == "Upload PDF":
            uploaded_file_br = st.file_uploader("Upload PDF of Funding Call", type=["pdf"], key="br_upload_funding_call")
            if uploaded_file_br:
                pdf_reader = PyPDF2.PdfReader(uploaded_file_br)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text()
                funding_call_text_br = text
                st.success("PDF uploaded and extracted successfully!")
        elif funding_call_option_br == "Enter URL":
            url_br = st.text_input("Enter URL of Funding Call", key="br_url_funding_call")
            if url_br:
                try:
                    downloaded = trafilatura.fetch_url(url_br)
                    if downloaded:
                        text = trafilatura.extract(downloaded, favor_recall=True)
                        funding_call_text_br = text if text else "Could not extract sufficient text from the URL."
                        st.success("Content extracted from URL!")
                    else:
                        st.warning("Could not fetch content from the provided URL.")
                except Exception as e:
                    st.error(f"Error fetching or parsing URL: {e}")
        
        # Extract and display funding call details
        if funding_call_text_br:
            extracted_fields_br = extract_fields(funding_call_text_br)
            st.session_state['funding_agency'] = extracted_fields_br.get('Funding Agency', '')
            st.session_state['scheme_type'] = extracted_fields_br.get('Scheme Type', '')
            st.session_state['thrust_areas'] = extracted_fields_br.get('Thrust Areas', '')
            st.session_state['eligibility'] = extracted_fields_br.get('Eligibility', '')
            st.session_state['duration'] = extracted_fields_br.get('Duration', '')
            st.session_state['budget'] = extracted_fields_br.get('Budget', '')
            st.session_state['submission_format'] = extracted_fields_br.get('Submission Format', '')
            st.session_state['last_date_submission'] = extracted_fields_br.get('Last Date of Submission', '')
    
    with col2:
        st.subheader("3. Proposal Content to Review")
        proposal_input_option = st.radio(
            "How would you like to provide your proposal?",
            ("Paste Text", "Upload PDF", "Upload DOCX"),
            key="br_proposal_input_option"
        )
        
        proposal_text_br = ""
        if proposal_input_option == "Paste Text":
            proposal_text_br = st.text_area("Paste Your Proposal Text Here", height=200, key="br_paste_proposal")
        elif proposal_input_option == "Upload PDF":
            uploaded_proposal_pdf = st.file_uploader("Upload PDF of Your Proposal", type=["pdf"], key="br_upload_proposal_pdf")
            if uploaded_proposal_pdf:
                pdf_reader = PyPDF2.PdfReader(uploaded_proposal_pdf)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text()
                proposal_text_br = text
                st.success("Proposal PDF uploaded and extracted successfully!")
        elif proposal_input_option == "Upload DOCX":
            uploaded_proposal_docx = st.file_uploader("Upload DOCX of Your Proposal", type=["docx"], key="br_upload_proposal_docx")
            if uploaded_proposal_docx:
                if docx is None:
                    st.error("DOCX support requires python-docx. Please upload PDF or paste text instead.")
                else:
                    doc = docx.Document(uploaded_proposal_docx)
                    proposal_text_br = "\n".join(p.text for p in doc.paragraphs)
                    st.success("Proposal DOCX uploaded and extracted successfully!")
        
        # Update session state with proposal content
        if proposal_text_br:
            st.session_state['full_proposal_draft'] = proposal_text_br
    
    st.markdown("---")
    
    # Researcher Profile Section
    st.subheader("4. Your Research Background")
    user_research_profile = st.text_area(
        "Describe your research background and key expertise (Publications, Patents, Grants, etc.)", 
        value=st.session_state.get('user_research_profile', ''), 
        height=150, 
        key="br_user_research_profile"
    )
    st.session_state['user_research_profile'] = user_research_profile
    
    # Template Sections Section
    st.subheader("5. Proposal Template Sections")
    template_sections_input_br = st.text_area(
        "Enter the template sections your proposal should follow (one per line)",
        value=st.session_state.get('actual_template_sections_used', ''),
        height=150,
        key="br_template_sections_input"
    )
    st.session_state['actual_template_sections_used'] = template_sections_input_br

    # Fetch data from session state (populated from inputs above)
    funding_agency = st.session_state.get('funding_agency', '')
    scheme_type = st.session_state.get('scheme_type', '')
    thrust_areas = st.session_state.get('thrust_areas', '')
    eligibility = st.session_state.get('eligibility', '')
    full_proposal_draft = st.session_state.get('full_proposal_draft', '')
    actual_template_sections_used = st.session_state.get('actual_template_sections_used', '')
    
    MAX_TEXT_LENGTH_FOR_PROMPT = 4000 # Max characters for summary before sending to AI

    # Summarize lengthy inputs to avoid exceeding token limits
    summarized_user_research_profile = summarize_text_for_prompt(user_research_profile, max_length=MAX_TEXT_LENGTH_FOR_PROMPT)
    summarized_full_proposal_draft = summarize_text_for_prompt(full_proposal_draft, max_length=MAX_TEXT_LENGTH_FOR_PROMPT)

    st.write(f"DEBUG: user_research_profile: {bool(user_research_profile)}")
    st.write(f"DEBUG: funding_agency: {bool(funding_agency)}")
    st.write(f"DEBUG: scheme_type: {bool(scheme_type)}")
    st.write(f"DEBUG: full_proposal_draft: {bool(full_proposal_draft)}")
    st.write(f"DEBUG: actual_template_sections_used: {bool(actual_template_sections_used)}")
    st.write(f"DEBUG (Brainstorm Room): full_proposal_draft content: '{full_proposal_draft[:50]}...' (length: {len(full_proposal_draft)})")
    st.write(f"DEBUG (Brainstorm Room): actual_template_sections_used content: '{actual_template_sections_used[:50]}...' (length: {len(actual_template_sections_used)})")

   
    st.markdown("**Proposal Draft Content (Generated/Used):**")
    if st.session_state.get('full_proposal_draft'):
        st.text_area("", value=full_proposal_draft, height=300, key="br_proposal_draft_display", disabled=True)
        if st.session_state.get('actual_template_sections_used'):
            st.markdown(f"**Based on Template Sections:**\n{actual_template_sections_used}")
        else:
            st.info("No template sections found for the generated proposal. Please ensure a template was used during proposal generation.")
    else:
        st.info("Please go to 'Grant Proposal Overview Generator' to generate a proposal draft.")
            
    # New: analysis mode toggle
    analysis_mode = st.radio(
        "Analysis Mode",
        ("Single-pass (recommended)", "Per-section (advanced)"),
        index=0,
        key="br_analysis_mode"
    )

    if st.button("Generate Brainstorm Analysis", key="generate_brainstorm_btn"):
        
        if user_research_profile and funding_agency and scheme_type and full_proposal_draft and actual_template_sections_used:
            if analysis_mode == "Single-pass (recommended)":
                # Show progress for preparation steps
                prep_progress = st.progress(0, text="Preparing data for analysis...")
                
                # Summarize user profile
                prep_progress.progress(33, text="Summarizing researcher profile to fit AI context window...")
                summarized_user_research_profile = summarize_text_for_prompt(user_research_profile, max_length=MAX_TEXT_LENGTH_FOR_PROMPT)
                
                # Summarize proposal
                prep_progress.progress(66, text="Summarizing proposal draft to fit AI context window...")
                summarized_full_proposal_draft = summarize_text_for_prompt(full_proposal_draft, max_length=MAX_TEXT_LENGTH_FOR_PROMPT)
                
                # Parse template sections
                prep_progress.progress(100, text="Preparation complete! Building analysis prompt...")
                template_section_titles = [line.strip() for line in actual_template_sections_used.split('\n') if line.strip()]
                template_sections_bullets = "\n".join(f"- {t}" for t in template_section_titles)

                single_pass_prompt = f"""
                You are a critical grant evaluator. Analyze the following research proposal draft against the funding call and researcher profile.

                --- Researcher Profile ---
                {summarized_user_research_profile}

                --- Funding Call Details ---
                Funding Agency: {funding_agency}
                Scheme Type: {scheme_type}
                Thrust Areas: {thrust_areas}
                Eligibility: {eligibility}

                --- Proposal Draft (summarized) ---
                {summarized_full_proposal_draft}

                --- Template Sections ---
                {template_sections_bullets}

                For EACH template section, output the following structure exactly:
                ### [Section Title]
                **Strengths**
                - ...
                - ...
                **Weaknesses**
                - ...
                - ...
                **Recommendations**
                - ...
                - ...

                Notes:
                - If the section content is missing or vague, state that clearly and recommend precise content to add.
                - Be specific and actionable. Avoid generic advice.
                """

                prep_progress.empty()  # Clear preparation progress bar
                with st.spinner("Generating brainstorm analysis (single-pass)..."):
                    single_resp = generate_content_with_retry(SELECTED_MODEL, single_pass_prompt)
                if single_resp and single_resp.text:
                    st.session_state['brainstorm_analysis_report'] = single_resp.text
                    st.success("Brainstorm analysis generated!")
                    st.rerun()
                else:
                    st.error("Failed to generate the analysis. Please try again.")
            else:
                # Existing per-section detailed analysis
                # Show progress for initial preparation
                prep_progress = st.progress(0, text="Preparing data for per-section analysis...")
                
                # Split proposal into sections
                prep_progress.progress(25, text="Splitting proposal into sections...")
                proposal_sections_content = split_proposal_into_sections(full_proposal_draft, actual_template_sections_used)
                all_section_reports = []

                # Summarize user profile once for all sections
                prep_progress.progress(50, text="Summarizing researcher profile to fit AI context window...")
                summarized_user_research_profile = summarize_text_for_prompt(user_research_profile, max_length=MAX_TEXT_LENGTH_FOR_PROMPT)
                
                prep_progress.progress(75, text="Parsing template sections...")
                template_section_titles = [line.strip() for line in actual_template_sections_used.split('\n') if line.strip()]

                prep_progress.progress(100, text="Preparation complete! Starting section-by-section analysis...")
                prep_progress.empty()  # Clear preparation progress bar
                
                total_sections = len(template_section_titles)
                progress_text = "Generating brainstorm analysis... Please wait."
                my_bar = st.progress(0, text=progress_text)
                
                for i, section_title_from_template in enumerate(template_section_titles):
                    section_content = proposal_sections_content.get(section_title_from_template, "").strip()
                    
                    # Show summarization progress for each section
                    my_bar.progress(
                        int((i / total_sections) * 100), 
                        text=f"Summarizing section '{section_title_from_template}' to fit AI context window..."
                    )
                    summarized_section_content = summarize_text_for_prompt(section_content, max_length=MAX_TEXT_LENGTH_FOR_PROMPT)

                    percent_complete = int(((i + 1) / total_sections) * 100)
                    my_bar.progress(percent_complete, text=f"Analyzing section: {section_title_from_template} ({percent_complete}% complete)")

                    section_brainstorm_prompt = f"""
                    --- Researcher Profile ---
                    {summarized_user_research_profile}

                    --- Funding Call Details ---
                    Funding Agency: {funding_agency}
                    Scheme Type: {scheme_type}
                    Thrust Areas: {thrust_areas}
                    Eligibility: {eligibility}

                    --- Specific Proposal Section for Analysis ---
                    Section Title: {section_title_from_template}
                    Section Content:
                    {summarized_section_content}

                    --- Instructions for AI ---
                    For the "{section_title_from_template}" section, provide a **highly critical and actionable** analysis. Focus on the following:
                    **{section_title_from_template}**
                    **Strengths**
                    - 2-3 strengths linked to call priorities or researcher's contributions
                    **Weaknesses**
                    - 2-3 critical weaknesses or gaps
                    **Recommendations**
                    - 2-3 specific actions to fix weaknesses and improve competitiveness
                    If the section content is empty or vague, say so and list exact content to add.
                    """

                    section_response = generate_content_with_retry(SELECTED_MODEL, section_brainstorm_prompt)
                    if section_response and section_response.text:
                        all_section_reports.append(f"### {section_title_from_template}\n{section_response.text}")
                    else:
                        all_section_reports.append(f"### {section_title_from_template}\n*Failed to generate analysis for this section.*")

                st.session_state['brainstorm_analysis_report'] = "\n\n---\n\n".join(all_section_reports)
                my_bar.progress(100, text="Brainstorm analysis complete!")
                st.success("Brainstorm analysis generated!")
                st.rerun()
        else:
            st.warning("""To generate brainstorm analysis, please ensure:
            1. Your Researcher Profile is provided in 'Align with My Research' or 'Grant Proposal Overview Generator'.
            2. Funding Call Details are loaded in 'Grant Proposal Overview Generator'.
            3. A Proposal Draft has been generated in 'Grant Proposal Overview Generator'.""")

    if 'brainstorm_analysis_report' in st.session_state:
        st.subheader("Brainstorm Analysis Report:")
        st.write(st.session_state['brainstorm_analysis_report'])
        st.download_button(
            label="Download Brainstorm Report",
            data=st.session_state['brainstorm_analysis_report'].encode('utf-8'),
            file_name="brainstorm_analysis_report.txt",
            mime="text/plain"
        )

elif st.session_state['current_main_view'] == 'draft_final':
    st.header("✍️ Draft Final Proposal")
    st.markdown("### Generate an Improved Proposal Based on Brainstorming Feedback")
    
    st.info("""
    This feature uses AI to draft an improved final proposal by:
    1. Taking your original proposal draft
    2. Analyzing the weaknesses identified in brainstorming
    3. Applying the recommendations provided
    4. Generating enhanced content for each section
    """)
    
    # --- Load Proposal Options ---
    saved_proposals = load_all_proposals()
    
    if not saved_proposals:
        st.warning("No saved proposals found. Please generate a proposal and run brainstorming analysis first.")
        st.markdown("""
        **Steps to use this feature:**
        1. Go to 'Grant Proposal Overview Generator' and generate a proposal
        2. Go to 'Brainstorm Room' and analyze your proposal
        3. Return here to draft the final improved version
        """)
    else:
        # Filter proposals that have brainstorm reports (handle None values)
        proposals_with_analysis = [p for p in saved_proposals if (p.get('brainstorm_analysis_report') or '').strip()]
        
        if not proposals_with_analysis:
            st.warning("No proposals with brainstorming analysis found. Please run brainstorming on your proposals first.")
        else:
            st.subheader("1. Select Proposal to Improve")
            
            proposal_options = []
            for idx, p in enumerate(proposals_with_analysis):
                title = f"{p['timestamp']} - {p['funding_agency']} - {p['scheme_type']}"
                proposal_options.append((title, idx))
            
            display_titles = [title for title, _ in proposal_options]
            selected_title = st.selectbox(
                "Choose a proposal with brainstorming feedback:", 
                display_titles, 
                key="draft_final_selector"
            )
            
            if selected_title:
                option_to_index = {title: idx for title, idx in proposal_options}
                original_idx = option_to_index[selected_title]
                selected_proposal = proposals_with_analysis[original_idx]
                
                # Display proposal details
                st.markdown("---")
                st.subheader("2. Proposal Details")
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.metric("Funding Agency", selected_proposal.get('funding_agency', 'N/A'))
                with col2:
                    st.metric("Scheme Type", selected_proposal.get('scheme_type', 'N/A'))
                with col3:
                    st.metric("Duration", selected_proposal.get('duration', 'N/A'))
                
                # Show original proposal in expandable section
                with st.expander("📄 View Original Proposal", expanded=False):
                    st.text_area(
                        "Original Proposal Content",
                        value=selected_proposal.get('full_proposal_content', 'No content'),
                        height=300,
                        disabled=True,
                        key="draft_final_original"
                    )
                
                # Show brainstorm analysis
                with st.expander("🧠 View Brainstorming Analysis", expanded=True):
                    st.markdown(selected_proposal.get('brainstorm_analysis_report', 'No analysis'))
                
                st.markdown("---")
                st.subheader("3. Generate Improved Final Proposal")
                
                # Improvement options
                improvement_approach = st.radio(
                    "Select improvement approach:",
                    ("Conservative (minor edits)", "Moderate (balanced improvements)", "Aggressive (major rewrite)"),
                    index=1,
                    key="improvement_approach"
                )
                
                if st.button("🚀 Generate Improved Final Proposal", key="generate_final_btn"):
                    # Extract data (handle None values from database)
                    original_proposal = selected_proposal.get('full_proposal_content') or ''
                    brainstorm_report = selected_proposal.get('brainstorm_analysis_report') or ''
                    template_sections = selected_proposal.get('template_sections') or ''
                    user_profile = selected_proposal.get('user_research_background') or ''
                    funding_agency = selected_proposal.get('funding_agency', '')
                    scheme_type = selected_proposal.get('scheme_type', '')
                    thrust_areas = selected_proposal.get('thrust_areas', '')
                    eligibility = selected_proposal.get('eligibility', '')
                    
                    if not all([original_proposal, brainstorm_report, template_sections]):
                        st.error("Missing required data. Please ensure the proposal has all necessary information.")
                    else:
                        # Define max text length for prompt
                        MAX_TEXT_LENGTH_FOR_PROMPT = 4000  # Max characters for summary before sending to AI
                        
                        # Show progress
                        progress_bar = st.progress(0, text="Preparing to generate improved proposal...")
                        
                        # Parse template sections
                        progress_bar.progress(10, text="Parsing proposal structure...")
                        template_section_titles = [line.strip() for line in template_sections.split('\n') if line.strip()]
                        
                        # Split original proposal into sections
                        progress_bar.progress(20, text="Extracting original sections...")
                        original_sections = split_proposal_into_sections(original_proposal, template_sections)
                        
                        # Summarize inputs
                        progress_bar.progress(30, text="Summarizing brainstorming feedback...")
                        summarized_brainstorm = summarize_text_for_prompt(brainstorm_report, max_length=MAX_TEXT_LENGTH_FOR_PROMPT)
                        summarized_user_profile = summarize_text_for_prompt(user_profile, max_length=MAX_TEXT_LENGTH_FOR_PROMPT)
                        
                        # Set improvement level instructions
                        if improvement_approach == "Conservative (minor edits)":
                            improvement_instruction = "Make minimal, targeted improvements. Keep the original structure and most content. Only address critical weaknesses."
                        elif improvement_approach == "Moderate (balanced improvements)":
                            improvement_instruction = "Make balanced improvements addressing all identified weaknesses while preserving strong elements. Enhance clarity and impact."
                        else:  # Aggressive
                            improvement_instruction = "Significantly rewrite sections to maximize impact. Address all weaknesses comprehensively and elevate the entire proposal quality."
                        
                        # Generate improved sections
                        improved_sections = {}
                        total_sections = len(template_section_titles)
                        
                        for i, section_title in enumerate(template_section_titles):
                            percent = int(30 + (i / total_sections) * 60)  # 30% to 90%
                            progress_bar.progress(
                                percent, 
                                text=f"Improving section '{section_title}' ({i+1}/{total_sections})..."
                            )
                            
                            original_section_content = original_sections.get(section_title, "").strip()
                            
                            # Build improvement prompt
                            improvement_prompt = f"""
You are an expert grant proposal writer. Your task is to improve the following proposal section based on brainstorming feedback.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
CONTEXT
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Funding Agency: {funding_agency}
Scheme Type: {scheme_type}
Thrust Areas: {thrust_areas}
Eligibility: {eligibility}

Researcher Profile (Summarized):
{summarized_user_profile}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
SECTION TO IMPROVE: "{section_title}"
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
ORIGINAL CONTENT:
{original_section_content if original_section_content else "[SECTION IS EMPTY OR MISSING]"}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
BRAINSTORMING FEEDBACK (Full Report):
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
{summarized_brainstorm}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
IMPROVEMENT INSTRUCTIONS
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Improvement Level: {improvement_instruction}

Your task:
1. Carefully review the original "{section_title}" section
2. Identify relevant weaknesses and recommendations from the brainstorming report for THIS section
3. Generate an IMPROVED version that:
   - Addresses all identified weaknesses
   - Implements the recommendations
   - Maintains alignment with funding agency priorities
   - Preserves the researcher's voice and authentic expertise
   - Uses clear, compelling, professional language
   - Includes specific details, metrics, and evidence where possible

CRITICAL RULES:
- Write ONLY the improved section content (no meta-commentary)
- Do NOT include section headers or labels
- Do NOT write "Here is the improved section" or similar phrases
- If the original section was empty, create comprehensive content based on recommendations
- Maintain appropriate length for this section type
- Ensure coherence with other proposal sections

OUTPUT: Write the improved section content directly below:
"""
                            
                            # Generate improved content
                            response = generate_content_with_retry(SELECTED_MODEL, improvement_prompt)
                            
                            if response and response.text:
                                improved_sections[section_title] = response.text.strip()
                            else:
                                improved_sections[section_title] = f"[Failed to improve this section. Original content retained.]\n\n{original_section_content}"
                        
                        # Combine improved sections
                        progress_bar.progress(95, text="Assembling final proposal...")
                        
                        improved_full_proposal = ""
                        for section_title in template_section_titles:
                            improved_full_proposal += f"\n\n## {section_title}\n\n"
                            improved_full_proposal += improved_sections.get(section_title, "[Content not generated]")
                        
                        improved_full_proposal = improved_full_proposal.strip()
                        
                        # Save to session state
                        st.session_state['improved_final_proposal'] = improved_full_proposal
                        st.session_state['improvement_metadata'] = {
                            'original_proposal_id': selected_proposal.get('id', 'N/A'),
                            'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                            'approach': improvement_approach,
                            'funding_agency': funding_agency,
                            'scheme_type': scheme_type
                        }
                        
                        progress_bar.progress(100, text="Improved proposal generated successfully!")
                        progress_bar.empty()
                        st.success("✅ Improved final proposal generated successfully!")
                        st.rerun()
                
                # Display improved proposal if it exists
                if 'improved_final_proposal' in st.session_state:
                    st.markdown("---")
                    st.subheader("4. Review Improved Final Proposal")
                    
                    # Show metadata
                    if 'improvement_metadata' in st.session_state:
                        metadata = st.session_state['improvement_metadata']
                        st.info(f"Generated: {metadata['timestamp']} | Approach: {metadata['approach']}")
                    
                    # Tabs for comparison
                    tab1, tab2, tab3 = st.tabs(["✨ Improved Proposal", "📊 Side-by-Side", "📄 Original"])
                    
                    with tab1:
                        st.markdown("### Improved Final Proposal")
                        improved_content = st.text_area(
                            "Edit if needed:",
                            value=st.session_state['improved_final_proposal'],
                            height=500,
                            key="improved_proposal_edit"
                        )
                        st.session_state['improved_final_proposal'] = improved_content
                    
                    with tab2:
                        st.markdown("### Side-by-Side Comparison")
                        col_orig, col_impr = st.columns(2)
                        with col_orig:
                            st.markdown("**Original**")
                            st.text_area(
                                "",
                                value=selected_proposal.get('full_proposal_content', ''),
                                height=400,
                                disabled=True,
                                key="comparison_original"
                            )
                        with col_impr:
                            st.markdown("**Improved**")
                            st.text_area(
                                "",
                                value=st.session_state['improved_final_proposal'],
                                height=400,
                                disabled=True,
                                key="comparison_improved"
                            )
                    
                    with tab3:
                        st.markdown("### Original Proposal")
                        st.text_area(
                            "",
                            value=selected_proposal.get('full_proposal_content', ''),
                            height=500,
                            disabled=True,
                            key="original_full_view"
                        )
                    
                    st.markdown("---")
                    st.subheader("5. Save or Export")
                    
                    col_save, col_export = st.columns(2)
                    
                    with col_save:
                        if st.button("💾 Save Improved Proposal to Database", key="save_improved_btn"):
                            # Save as new proposal
                            new_proposal_data = {
                                'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                                'funding_agency': selected_proposal.get('funding_agency', ''),
                                'scheme_type': selected_proposal.get('scheme_type', '') + " (FINAL DRAFT)",
                                'duration': selected_proposal.get('duration', ''),
                                'budget': selected_proposal.get('budget', ''),
                                'thrust_areas': selected_proposal.get('thrust_areas', ''),
                                'eligibility': selected_proposal.get('eligibility', ''),
                                'submission_format': selected_proposal.get('submission_format', ''),
                                'user_research_background': selected_proposal.get('user_research_background', ''),
                                'template_sections': selected_proposal.get('template_sections', ''),
                                'full_proposal_content': st.session_state['improved_final_proposal'],
                                'brainstorm_analysis_report': f"IMPROVED VERSION - Original analysis:\n\n{selected_proposal.get('brainstorm_analysis_report') or ''}",
                                'alignment_score': selected_proposal.get('alignment_score', None)
                            }
                            save_proposal_to_db(new_proposal_data)
                            st.success("✅ Improved proposal saved as new draft!")
                    
                    with col_export:
                        # Export buttons
                        st.download_button(
                            label="📥 Download as TXT",
                            data=st.session_state['improved_final_proposal'].encode('utf-8'),
                            file_name=f"final_proposal_{selected_proposal.get('scheme_type', 'draft').replace(' ', '_')}.txt",
                            mime="text/plain"
                        )

elif st.session_state['current_main_view'] == 'my_drafts':
    st.header("📝 My Drafts & Submissions")
    st.markdown("### Review and Manage Your Saved Proposal Drafts")

    saved_proposals = load_all_proposals()

    if saved_proposals:
        proposal_options = []
        for idx, p in enumerate(saved_proposals):
            title = f"{p['timestamp']} - {p['funding_agency']} - {p['scheme_type']}"
            proposal_options.append((title, idx)) # Store (display_text, original_index)

        option_to_index = {title: idx for title, idx in proposal_options}
        display_titles = [title for title, _ in proposal_options]

        selected_title = st.selectbox("Select a saved proposal to review:", display_titles, key="my_drafts_proposal_selector")

        if selected_title:
            original_idx = option_to_index[selected_title]
            selected_proposal_data = saved_proposals[original_idx]
            
            st.markdown(f"**Selected Proposal Details:**")
            st.json({
                "Funding Agency": selected_proposal_data.get('funding_agency', 'N/A'),
                "Scheme Type": selected_proposal_data.get('scheme_type', 'N/A'),
                "Timestamp": selected_proposal_data.get('timestamp', 'N/A')
            })

            st.subheader("Full Proposal Content:")
            st.text_area("", value=selected_proposal_data.get('full_proposal_content', 'No content available.'), height=500, disabled=True, key="my_drafts_full_proposal")
            
            brainstorm_report = selected_proposal_data.get('brainstorm_analysis_report', '')
            if brainstorm_report:
                st.subheader("Brainstorm Analysis Report (Editable):")
                
                parsed_report = parse_brainstorm_report(brainstorm_report)
                
                # Display editable sections
                edited_strengths = st.text_area("**Strengths:**", value=parsed_report.get('Strengths', ''), height=150, key="br_strengths_edit")
                edited_weaknesses = st.text_area("**Weaknesses:**", value=parsed_report.get('Weaknesses', ''), height=150, key="br_weaknesses_edit")
                edited_recommendations = st.text_area("**Recommendations:**", value=parsed_report.get('Recommendations', ''), height=150, key="br_recommendations_edit")
                
                # You could add a button here to save the edited brainstorm report back to the database
                # For now, it's just editable in the session.

            # Add a button to load this proposal into the generator for further editing
            if st.button("Load this Draft into Proposal Generator for Editing", key="my_drafts_load_for_edit"):
                st.session_state['funding_agency'] = selected_proposal_data.get('funding_agency', '')
                st.session_state['scheme_type'] = selected_proposal_data.get('scheme_type', '')
                st.session_state['duration'] = selected_proposal_data.get('duration', '')
                st.session_state['budget'] = selected_proposal_data.get('budget', '')
                st.session_state['thrust_areas'] = selected_proposal_data.get('thrust_areas', '')
                st.session_state['eligibility'] = selected_proposal_data.get('eligibility', '')
                st.session_state['submission_format'] = selected_proposal_data.get('submission_format', '')
                st.session_state['user_research_profile'] = selected_proposal_data.get('user_research_background', '')
                st.session_state['actual_template_sections_used'] = selected_proposal_data.get('template_sections', '')
                st.session_state['full_proposal_draft'] = selected_proposal_data.get('full_proposal_content', '')
                st.session_state['brainstorm_analysis_report'] = selected_proposal_data.get('brainstorm_analysis_report', '')
                
                st.success(f"Draft '{selected_title}' loaded into Proposal Generator!")
                nav_to('proposal_generator') # Navigate to the generator page
                st.rerun()
    else:
        st.info("No proposal drafts saved yet. Generate and save a proposal to see it here!")

elif st.session_state['current_main_view'] == 'export_share':
    st.header("📤 Export & Share Center")
